#!/usr/bin/env python3
"""Generate GTM Report v1 as DOCX for NapoleonIT — client-facing."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)

# --- Helpers ---
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        set_cell_shading(cell, 'F0F4F8')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_bold_paragraph(doc, bold_text, normal_text=''):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_bullet(doc, bold_part, text):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part:
        r = p.add_run(bold_part)
        r.bold = True
    if text:
        p.add_run(text)
    return p

def add_quote_block(doc, text):
    """Add an indented quote-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
    r.font.italic = True
    return p

def add_note(doc, text):
    """Add a note/caveat paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    r.font.italic = True
    return p

def add_sequence_line(doc, day, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(day)
    r.bold = True
    r.font.size = Pt(10)
    p.add_run('  ' + text).font.size = Pt(10)
    return p


# ========== DOCUMENT ==========

# --- Title page ---
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading('Стратегия выхода на рынок', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Система «ТОР»')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Napoleon IT')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta2.add_run('Подготовлено: FirstLeads  |  Февраль 2026')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_page_break()

# ===== О ДОКУМЕНТЕ =====
doc.add_heading('О документе', level=1)

p = doc.add_paragraph(
    'Это стратегия первой кампании по выводу системы «ТОР» на рынок. '
    'Документ описывает: как мы поняли продукт, на какой сегмент нацелена первая кампания, '
    'что именно мы будем говорить и делать, какие результаты ожидаем.'
)

p = doc.add_paragraph(
    'Цель документа — согласовать с вами подход до начала работы. '
    'Нам важно, чтобы вы подтвердили: мы правильно поняли продукт, '
    'идём к правильным людям и говорим правильные вещи.'
)

add_bold_paragraph(doc, 'Что направим отдельно на согласование:')
add_bullet(doc, 'Полный список целевых компаний ', '— с приоритетами и обоснованием выбора каждой')
add_bullet(doc, 'Точные формулировки писем ', '— все варианты для A/B тестирования, финальные тексты')
add_bullet(doc, 'Скрипты телефонных звонков ', '— полные сценарии с обработкой возражений')
add_bullet(doc, 'Шаблоны Telegram-сообщений ', '— для контактов из профессиональных сообществ')

p = doc.add_paragraph(
    'Всё, что уходит от имени вашего продукта, сначала проходит через вас.'
)

doc.add_paragraph()

# ===== 1. КАК МЫ ПОНЯЛИ ВАШ ПРОДУКТ =====
doc.add_heading('1. Как мы поняли ваш продукт', level=1)

p = doc.add_paragraph(
    'Система «ТОР» — платформа видеофиксации обходов на объектах. '
    'Инспектор идёт по объекту с камерой, а система автоматически строит его маршрут на плане '
    'с точностью до 10 см и привязывает каждую точку маршрута к видеозаписи. '
    'Работает внутри зданий и на открытых площадках, где нет GPS (глушилки, подземные конструкции). '
    'Не требует никакой инфраструктуры — ни маяков, ни меток, ни монтажа. '
    'Запуск на новом объекте занимает часы, а не недели.'
)

p = doc.add_paragraph(
    'Продукт уже работает на пилоте у АПРИ (ЖК «Грани», Челябинск) в сегменте стройконтроля. '
    'Разработан совместно Napoleon IT и НИИСТРОМ.'
)

add_bold_paragraph(doc, 'Ключевое отличие от альтернатив:')
add_bullet(doc, 'Против UWB-маяков (Navigine, RTLS): ', 'не нужна инфраструктура — экономия недель на монтаж и сотни тысяч рублей на оборудование')
add_bullet(doc, 'Против NFC/QR-меток (VGL, «Стража»): ', 'невозможно подделать — система записывает реальный маршрут с видео, а не контакт с меткой')
add_bullet(doc, 'Против зарубежных CV-платформ (OpenSpace, Buildots): ', 'единственное решение на российском рынке после санкционного ухода конкурентов')

doc.add_paragraph()

# ===== 2. СЕГМЕНТЫ РЫНКА =====
doc.add_heading('2. Сегменты рынка', level=1)

p = doc.add_paragraph(
    'Мы исследовали 9 потенциальных сегментов. Три из них — перспективные. '
    'Начинаем с одного — стройконтроля генподрядчиков, где уже есть пилот и понятный покупатель. '
    'Остальные два — на следующем этапе.'
)

doc.add_heading('Куда идём сейчас: Стройконтроль генподрядчиков', level=2)

add_bold_paragraph(doc, 'Кто эти компании:')
p = doc.add_paragraph(
    'Крупные генподрядные компании в жилищном строительстве — выручка от 1 млрд \u20bd в год, '
    '5 и более объектов одновременно. Именно на таком масштабе контроль обходов становится критичным, '
    'а споры с субподрядчиками — регулярными.'
)

add_bold_paragraph(doc, 'Кто принимает решение:')
p = doc.add_paragraph(
    'Начальник стройконтроля или директор по строительству. Этот человек ежедневно сталкивается '
    'с ситуацией «слово против слова» — инспектор говорит, что обошёл объект, подрядчик утверждает '
    'обратное. Разбирательства тянутся неделями, а доказательств нет.'
)

add_bold_paragraph(doc, 'Почему начинаем с них:')
add_bullet(doc, None, 'Уже есть работающий пилот в этом сегменте (АПРИ, Челябинск)')
add_bullet(doc, None, 'Боль предварительно подтверждена действующим пилотом')
add_bullet(doc, None, 'Регуляторное давление: СП 543 обязывает фиксировать стройконтроль электронно; ТИМ обязателен с 2025 года, но ~40% площадей до сих пор строятся без него')

add_bold_paragraph(doc, 'Примеры компаний:')
add_table(doc,
    ['Компания', 'Почему подходит'],
    [
        ['ООО «Монолит» (Краснодар)', '40.5 млрд \u20bd, рост +50% — взрывное расширение создаёт нагрузку на контроль качества'],
        ['Брусника (Екатеринбург)', '#1 по качеству в рейтинге ЕРЗ — естественный покупатель инструментов контроля'],
        ['МСУ-1 / ГК ФСК (Москва)', 'Строят для Фонда реновации. Москва = зоны подавления GPS, прямой сценарий для ТОР'],
        ['Setl Group (Санкт-Петербург)', 'Уже инвестируют в лазерное сканирование — готовы к новым решениям'],
        ['ГК ТОЧНО', 'Используют платформу «Базис», есть руководитель цифровых продуктов'],
    ],
    col_widths=[5, 13]
)

add_bold_paragraph(doc, 'Что мы хотим проверить:')
add_bullet(doc, '1. ', 'Готовы ли начальники стройконтроля тестировать новый инструмент, имея Excel и WhatsApp «по привычке»?')
add_bullet(doc, '2. ', 'Достаточно ли одного пилота (АПРИ) как доказательства, или нужно больше кейсов?')

doc.add_paragraph()

# --- Куда пойдём дальше ---
doc.add_heading('Куда пойдём дальше', level=2)

p = doc.add_paragraph(
    'По результатам исследования мы видим ещё два перспективных направления. '
    'В них мы не идём прямо сейчас, но они логически вытекают из первого сегмента.'
)

add_bold_paragraph(doc, 'Приёмка квартир у девелоперов. ',
    'Мораторий на неустойки отменён с 01.01.2026 — прогноз: до 100 000 исков на 200 млрд \u20bd. '
    'Все платформы приёмки (Базис, Домиленд) работают на фото, а фото в суде оспаривают. '
    'Видеофиксацию при приёмке не предлагает никто. '
    'Важно: тот же девелопер, который купит ТОР для стройконтроля, — потенциальный покупатель и для приёмки. '
    'Однако у нас пока нет пилота в приёмке и не подтверждено, что видео даёт преимущество перед фото в суде.'
)

add_bold_paragraph(doc, 'Нефтегаз (НПЗ). ',
    '32 целевых нефтеперерабатывающих завода, обязательные регламентные обходы (ФЗ-116), '
    'NFC-метки можно обмануть. Крупные бюджеты. Но НПЗ — объекты критической инфраструктуры: '
    'продукт должен работать на серверах заказчика, а опыта такого развёртывания пока нет. '
    'Этот сегмент заблокирован до технической валидации серверной версии.'
)

doc.add_page_break()

# ===== 3. ЧТО МЫ БУДЕМ ДЕЛАТЬ =====
doc.add_heading('3. Что мы будем делать', level=1)

p = doc.add_paragraph()
r = p.add_run('Первая кампания — стройконтроль генподрядчиков. 60 контактов, 2 недели.')
r.bold = True

doc.add_heading('Оффер', level=2)

add_bold_paragraph(doc, 'В одном предложении:')
p = doc.add_paragraph(
    'Каждый спор с подрядчиком решается за 5 минут по видеозаписи обхода — '
    'запуск на объекте за 2 часа, без маяков, бесплатный пилот.'
)

add_bold_paragraph(doc, 'Почему это работает:')
p = doc.add_paragraph(
    'Один арбитражный спор с подрядчиком — от 500 тыс. \u20bd и недели нервов. '
    'ТОР стоит 25–40 тыс. \u20bd в месяц на объект — окупается при одном предотвращённом споре в год. '
    'СП 543 уже обязывает фиксировать стройконтроль электронно. '
    'Пилот АПРИ подтверждает, что решение работает. '
    'Бесплатный пилот на объекте клиента снижает риск до нуля.'
)

doc.add_heading('Как это выглядит на практике', level=2)

add_bold_paragraph(doc, 'Пример письма:')

add_quote_block(doc,
    '[Имя], увидел, что [ЗАЦЕПКА].\n\n'
    'Застройщики вашего масштаба часто сталкиваются с тем, что при спорах с подрядчиками '
    'нет доказательной базы — слово против слова, разбирательства на недели.\n\n'
    'Для АПРИ на ЖК «Грани» мы создали ИИ-видеофиксацию обходов: маршрут инспектора с видео '
    'появляется на плане здания с точностью до 10 см. Без маяков, из оборудования — каска с камерой, '
    'разворачивается за 2 часа.\n\n'
    'Актуально ли это для ваших объектов? Готов показать как это работает — за 15 минут в Zoom.'
)

add_note(doc, 'Это один из вариантов, которые мы будем тестировать. '
    'Альтернативные варианты подготовлены для A/B тестирования и будут направлены дополнительно.')

doc.add_paragraph()
add_bold_paragraph(doc, 'Пример звонка:')

add_quote_block(doc,
    '[Имя], здравствуйте, у вас найдется 3 минуты? [Звонящий], Napoleon IT.\n\n'
    'Я звоню, потому что видел, что [Компания] сейчас [персонализированная вставка]. Верно?\n\n'
    'Для похожей компании мы создали систему видеофиксации обходов стройконтроля. '
    'Она фиксирует маршрут инспектора на плане здания с точностью до 10 см. '
    'Без маяков, разворачивается за 2 часа, из оборудования нужна только каска с камерой.\n\n'
    'Уже работает у АПРИ на ЖК «Грани» в Челябинске, помогает тщательно отслеживать и фиксировать все обходы.\n\n'
    'Мы сейчас запускаем бесплатные пилоты для генподрядчиков. '
    'Хочу предложить вашей компании. Вам бы было это интересно?'
)

doc.add_paragraph()

# --- Как ищем контакты ---
doc.add_heading('Как ищем контакты', level=2)

p = doc.add_paragraph(
    'Целевой руководитель — начальник стройконтроля или директор по строительству. '
    'Компании с выручкой от 1 млрд \u20bd, 5+ объектов одновременно.'
)

p = doc.add_paragraph(
    'Контакты собираем через специализированные базы данных и открытые источники. '
    'Каждого руководителя исследуем индивидуально: сайт компании, вакансии, новости, '
    'профессиональный профиль. На каждый контакт — 12–18 минут подготовки. '
    'Ни одно письмо и ни один звонок не уходит без персональной привязки к ситуации конкретной компании.'
)

add_bold_paragraph(doc, 'Приоритетные цели первой волны: ',
    'Монолит (40.5 млрд \u20bd, +50% рост), Брусника (#1 по качеству ЕРЗ), '
    'МСУ-1/ГК ФСК (Москва, GPS-глушилки), Setl Group (СПб), ГК ТОЧНО.'
)

# --- Последовательность ---
doc.add_heading('Последовательность обращений', level=2)

add_sequence_line(doc, 'День 1', 'Персональное письмо')
add_sequence_line(doc, 'День 3', 'Звонок')
add_sequence_line(doc, 'День 5', 'Второе письмо (другой ракурс: стоимость одного спора + СП 543)')
add_sequence_line(doc, 'День 7', 'Второй звонок')
add_sequence_line(doc, 'День 10', 'Финальное письмо')

doc.add_paragraph()

p = doc.add_paragraph(
    'Каждое следующее касание меняет ракурс, а не повторяет предыдущее: '
    'результат \u2192 стоимость бездействия \u2192 конкретный кейс АПРИ \u2192 завершение.'
)

# --- Объём ---
doc.add_heading('Объём', level=2)

add_table(doc,
    ['Параметр', 'Значение'],
    [
        ['Контактов в кампании', '60 (30 + 30 для A/B теста)'],
        ['Подготовка контактов', '2–3 рабочих дня'],
        ['Длительность кампании', '2 недели'],
        ['Писем (всего)', '~240 (до 4 на контакт)'],
        ['Звонков (всего)', '~80–120'],
        ['Первые ответы', 'К концу первой недели'],
        ['Квалифицированные лиды', 'Начиная со второй недели'],
    ],
    col_widths=[6, 12]
)

# --- Когда меняем подход ---
doc.add_heading('Когда меняем подход', level=2)

add_bullet(doc, None, 'Письма открывают, но не отвечают \u2192 меняем ракурс оффера')
add_bullet(doc, None, 'Не открывают \u2192 меняем тему письма, проверяем доставляемость')
add_bullet(doc, None, 'Более 5% ответов «не интересно» \u2192 анализируем причины, возможно боль слабее, чем мы думали')
add_bullet(doc, None, 'Ноль ответов после 100 писем \u2192 останавливаем, пересматриваем оффер или сегмент')
add_bullet(doc, None, 'Работает хорошо \u2192 масштабируем на оставшиеся компании из списка')

doc.add_page_break()

# ===== 4. ЧТО ВЫ ПОЛУЧАЕТЕ =====
doc.add_heading('4. Что вы получаете', level=1)

doc.add_heading('По каждому лиду', level=2)

p = doc.add_paragraph('Каждый квалифицированный лид передаётся вам в виде структурированного отчёта:')

add_bullet(doc, 'Контактные данные ', '— имя, должность, компания, email, телефон')
add_bullet(doc, 'Подтверждённая проблема ', '— словами самого руководителя, не наша гипотеза')
add_bullet(doc, 'Текущее решение ', '— что используют, что не устраивает')
add_bullet(doc, 'Бюджет и сроки ', '— подтверждены или нет, есть ли активный проект для пилота')
add_bullet(doc, 'Рекомендация для демо ', '— что показать, чего избегать')
add_bullet(doc, 'История касаний ', '— все письма и звонки с результатами')

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Горячий лид ')
r.bold = True
p.add_run('передаётся в течение 24 часов. ')
r2 = p.add_run('Тёплый ')
r2.bold = True
p.add_run('— в течение 48 часов.')

# --- Еженедельный отчёт ---
doc.add_heading('Еженедельный отчёт', level=2)

p = doc.add_paragraph('Каждую пятницу:')

add_bullet(doc, None, 'Сколько контактов обработано, сколько ответов, звонков, встреч')
add_bullet(doc, None, 'Что работает, что нет — конкретно, с примерами')
add_bullet(doc, 'Обратная связь от рынка ', '— какие гипотезы подтвердились, какие нет, что узнали нового')
add_bullet(doc, None, 'Результаты A/B теста тем письма')
add_bullet(doc, None, 'План на следующую неделю')

doc.add_paragraph()

# ===== 5. ЧТО НУЖНО ОТ ВАС =====
doc.add_heading('5. Что нужно от вас', level=1)

items = [
    'Согласовать, что начинаем со стройконтроля генподрядчиков (раздел 2)',
    'Утвердить тон и содержание обращений — это пойдёт от имени вашего продукта (раздел 3)',
    'Есть ли кейсы или цифры, которые мы можем использовать в обращениях? (сэкономленное время, предотвращённые споры, отзывы инспекторов)',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_paragraph()

# ===== 6. ОГРАНИЧЕНИЯ И ДОПУЩЕНИЯ =====
doc.add_heading('6. Ограничения и допущения', level=1)

doc.add_heading('Что основано на данных', level=2)

facts = [
    'Конкурентная карта: 15+ решений на 4 уровнях, подтверждено исследованием',
    'Размер рынка: ~2 953 застройщика, ~102 млн м\u00b2 жилья (данные ЕРЗ.РФ)',
    'Регуляторные требования: СП 543, обязательность ТИМ — публичные документы',
    'Пилот АПРИ: подтверждён из нескольких источников (IT-World, Smart-Lab, Napoleon IT)',
    '68+ целевых компаний: идентифицированы поимённо',
]
for f in facts:
    add_bullet(doc, None, f)

doc.add_heading('Что основано на гипотезах', level=2)

add_bullet(doc, 'Гипотеза 1: ', 'Начальники стройконтроля готовы тестировать новый инструмент (даже при работающих Excel, фотофиксациях и WhatsApp). Проверяем outreach-ем.')
add_bullet(doc, 'Гипотеза 2: ', 'Одного пилота (АПРИ) достаточно как доказательства для первых продаж. Проверяем по конверсии в звонки.')
add_bullet(doc, 'Гипотеза 3: ', '25–40 тыс. \u20bd/мес на объект — приемлемая цена для генподрядчика. Проверяем по возражениям в звонках.')

doc.add_heading('Что может измениться после первой недели', level=2)

add_bullet(doc, None, 'Ракурс оффера: A/B тест покажет, что резонирует — результат (5 минут вместо 5 недель) или стоимость бездействия (500 тыс. за спор)')
add_bullet(doc, None, 'Целевой руководитель: если начальник стройконтроля не отвечает, но директор по строительству — да, переключаемся')
add_bullet(doc, None, 'Формулировка ценности: конкретные слова, которые использует рынок в ответах, заменят наши предположения')
add_bullet(doc, None, 'Если стройконтроль реагирует слабо — пересматриваем оффер или переключаемся на следующий сегмент (приёмка квартир) по согласованию с вами')

# ===== FOOTER =====
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Стратегия выхода на рынок  |  Система «ТОР»  |  Napoleon IT')
r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
r.font.size = Pt(9)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Подготовлено: FirstLeads  |  Февраль 2026')
r2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
r2.font.size = Pt(9)

# --- Save ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'gtm-report-v1.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
