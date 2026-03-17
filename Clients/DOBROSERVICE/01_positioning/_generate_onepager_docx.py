#!/usr/bin/env python3
"""Generate one-pager DOCX with segments table and elevator pitches for ДОБРОСЕРВИС Digital."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Colors
BLUE = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x05, 0x96, 0x69)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
GRAY_600 = RGBColor(0x4B, 0x55, 0x63)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_800 = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BG_BLUE_LIGHT = "DBEAFE"
BG_GREEN_LIGHT = "D1FAE5"
BG_GREEN_LIGHTER = "ECFDF5"
BG_YELLOW_LIGHT = "FEF3C7"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, side, color="000000", sz="4", val="single"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    border = parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
    )
    existing = tcBorders.find(qn(f'w:{side}'))
    if existing is not None:
        tcBorders.remove(existing)
    tcBorders.append(border)


def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = GRAY_800

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ── Title ──
title = doc.add_heading('ДОБРОСЕРВИС Digital — One-Pager', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Резюме позиционирования · 12 марта 2026')
r.font.color.rgb = GRAY_600
r.font.size = Pt(11)
r.font.name = 'Calibri'

doc.add_paragraph()

# ════════════════════════════════════════
# SEGMENTS TABLE
# ════════════════════════════════════════
h = doc.add_heading('Сегменты', level=1)
for run in h.runs:
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'

headers = ['Сегмент', 'ЛПР', 'Приоритет', 'Что подтверждено', 'Что проверяем']
rows = [
    ['Промышленность Свердловской обл., 500-1500 чел.', 'PR/GR-директор', 'Приоритет 1',
     '40-60 компаний, 0 конкурентов, GR-повод, цена под тендер', 'PR/GR быстрее принимает решение чем HR'],
    ['Промышленность Свердловской обл., 500-1500 чел.', 'HR-директор', 'Приоритет 1',
     'Дефицит 100K рабочих, травматизм +14.7%, 59% расширяют well-being', 'HR готов тратить при урезанном бюджете'],
    ['Строительство Свердловской обл., 200-1000 чел.', 'Генеральный', 'Приоритет 2',
     'Травматизм, низкая квалификация', 'ГД готов тратить на соцпрограмму'],
]

table = doc.add_table(rows=1 + len(rows), cols=len(headers))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h_text in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h_text
    set_cell_shading(cell, BG_BLUE_LIGHT)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = GRAY_700

for r_idx, row in enumerate(rows):
    for c_idx, val in enumerate(row):
        cell = table.cell(r_idx + 1, c_idx)
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
        # Highlight P1 rows
        if r_idx < 2:
            set_cell_shading(cell, BG_GREEN_LIGHTER)

# Set column widths
widths = [4.0, 2.5, 2.0, 4.5, 4.0]
for i, w in enumerate(widths):
    for row in table.rows:
        row.cells[i].width = Cm(w)

doc.add_paragraph()

# Strategy box
strat_table = doc.add_table(rows=1, cols=2)
strat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

left = strat_table.cell(0, 0)
left.width = Cm(0.3)
set_cell_shading(left, "059669")
left.text = ""

right = strat_table.cell(0, 1)
set_cell_shading(right, BG_GREEN_LIGHT)
p = right.paragraphs[0]
r = p.add_run('Стратегия: A/B-тест на одних и тех же компаниях')
r.font.bold = True
r.font.color.rgb = GREEN
r.font.name = 'Calibri'
r.font.size = Pt(11)
p2 = right.add_paragraph()
r2 = p2.add_run('Те же самые предприятия, два разных ЛПР. Цель: первые ответы, discovery-звонки, приглашения на конференцию 24.03.')
r2.font.name = 'Calibri'
r2.font.size = Pt(10)

for cell in [left, right]:
    for side in ['top', 'bottom', 'end', 'start']:
        set_cell_border(cell, side, val="none", sz="0")

doc.add_paragraph()
add_separator(doc)
doc.add_paragraph()

# ════════════════════════════════════════
# ELEVATOR PITCHES
# ════════════════════════════════════════
h = doc.add_heading('Elevator Pitches', level=1)
for run in h.runs:
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'

# PR/GR one-liner
p = doc.add_paragraph()
r = p.add_run('PR/GR — однострочник')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_600
r.font.name = 'Calibri'
r.font.bold = True

pitch_table = doc.add_table(rows=1, cols=2)
pitch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
left = pitch_table.cell(0, 0)
left.width = Cm(0.3)
set_cell_shading(left, "2563EB")
left.text = ""
right = pitch_table.cell(0, 1)
set_cell_shading(right, BG_BLUE_LIGHT)
p = right.paragraphs[0]
r = p.add_run('Готовый GR-инфоповод и соцпрограмма для сотрудников — подключение за дни, от 150 тыс./год.')
r.font.italic = True
r.font.name = 'Calibri'
r.font.size = Pt(11)
for cell in [left, right]:
    for side in ['top', 'bottom', 'end', 'start']:
        set_cell_border(cell, side, val="none", sz="0")

doc.add_paragraph()

# HR one-liner
p = doc.add_paragraph()
r = p.add_run('HR — однострочник')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_600
r.font.name = 'Calibri'
r.font.bold = True

pitch_table2 = doc.add_table(rows=1, cols=2)
pitch_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
left = pitch_table2.cell(0, 0)
left.width = Cm(0.3)
set_cell_shading(left, "2563EB")
left.text = ""
right = pitch_table2.cell(0, 1)
set_cell_shading(right, BG_BLUE_LIGHT)
p = right.paragraphs[0]
r = p.add_run('Цифровая программа поддержки сотрудников: 9 направлений, от 300 руб./чел./год, без тендера.')
r.font.italic = True
r.font.name = 'Calibri'
r.font.size = Pt(11)
for cell in [left, right]:
    for side in ['top', 'bottom', 'end', 'start']:
        set_cell_border(cell, side, val="none", sz="0")

doc.add_paragraph()

# Extended PR/GR pitch
p = doc.add_paragraph()
r = p.add_run('Расширенный (PR/GR)')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_600
r.font.name = 'Calibri'
r.font.bold = True

pitch_table3 = doc.add_table(rows=1, cols=2)
pitch_table3.alignment = WD_TABLE_ALIGNMENT.CENTER
left = pitch_table3.cell(0, 0)
left.width = Cm(0.3)
set_cell_shading(left, "2563EB")
left.text = ""
right = pitch_table3.cell(0, 1)
set_cell_shading(right, BG_BLUE_LIGHT)
p = right.paragraphs[0]
r = p.add_run(
    'Промышленные предприятия на 500-1500 человек в регионах тратят на ДМС миллионы, '
    'но у сотрудников нет поддержки по юридическим, финансовым и психологическим вопросам. '
    'Собственную программу разработать — дорого и долго. Полноценные EAP (типа «Понимаю») '
    'стоят от 600 тыс. и требуют тендера.'
)
r.font.italic = True
r.font.name = 'Calibri'
r.font.size = Pt(10)

p2 = right.add_paragraph()
p2.space_before = Pt(8)
r2 = p2.add_run(
    'ДОБРОСЕРВИС Digital — полностью цифровая программа поддержки за 150-500 тыс./год. '
    '9 направлений: робот-юрист, психолог-бот, финансовый консультант, гайды по жизненным ситуациям. '
    'Основана на 19-летней базе знаний ЕЮС — тех самых специалистов, которые обслуживают '
    'Северсталь, Газпромбанк и Ленту. Подключение за дни, без тендера. '
    'Бонус — участие в госпрограмме повышения качества труда с признанием от губернатора.'
)
r2.font.italic = True
r2.font.name = 'Calibri'
r2.font.size = Pt(10)

for cell in [left, right]:
    for side in ['top', 'bottom', 'end', 'start']:
        set_cell_border(cell, side, val="none", sz="0")

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('ДОБРОСЕРВИС Digital · FirstLeads · 12 марта 2026')
r.font.color.rgb = GRAY_600
r.font.size = Pt(9)
r.font.name = 'Calibri'

output = '/Users/wladyslaw/Documents/Job/FirstLeadsV2/Clients/DOBROSERVICE/01_positioning/onepager-segments-pitches.docx'
doc.save(output)
print(f'Saved: {output}')
