from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Вопросы для Андрея', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('ДОБРОСЕРВИС Digital · 12 марта 2026')
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
run.font.size = Pt(12)

doc.add_paragraph()

# Intro
intro = doc.add_paragraph()
intro_run = intro.add_run('Андрей, ниже 6 вопросов, без которых мы не можем начать аутрич. Первые два — блокеры (без них не сможем позвонить). Остальные можно решить параллельно. Ответь сам или передай кому нужно. Можно коротко — одним предложением.')
intro_run.font.size = Pt(11)

doc.add_paragraph()

# Helper
def add_question(num, priority, priority_color, title_text, context, questions, note=None):
    # Priority + number
    p = doc.add_paragraph()
    run_num = p.add_run(f'Вопрос {num}  ')
    run_num.bold = True
    run_num.font.size = Pt(14)

    run_pri = p.add_run(f'[{priority}]')
    run_pri.bold = True
    run_pri.font.size = Pt(11)
    run_pri.font.color.rgb = priority_color

    # Title
    h = doc.add_heading(title_text, level=2)

    # Context
    ctx = doc.add_paragraph()
    ctx_run = ctx.add_run(context)
    ctx_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    ctx_run.font.size = Pt(10)

    # Questions
    for q in questions:
        bullet = doc.add_paragraph(q, style='List Bullet')

    # Note
    if note:
        n = doc.add_paragraph()
        n_run = n.add_run(note)
        n_run.italic = True
        n_run.font.size = Pt(10)
        n_run.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C) if 'ASAP' in note else RGBColor(0x97, 0x5A, 0x16)

    # Answer area
    doc.add_paragraph()
    ans = doc.add_paragraph()
    ans_run = ans.add_run('Ответ:')
    ans_run.italic = True
    ans_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph()
    # Separator
    sep = doc.add_paragraph()
    sep_run = sep.add_run('─' * 60)
    sep_run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    sep_run.font.size = Pt(8)
    doc.add_paragraph()


RED = RGBColor(0xDC, 0x26, 0x26)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
GREEN = RGBColor(0x05, 0x96, 0x69)

add_question(
    1, 'БЛОКЕР — ASAP', RED,
    'Что конкретно входит в облегчённую версию?',
    'Когда позвоним в компанию, первый вопрос будет «а что конкретно получат наши сотрудники?». Сейчас мы знаем что «~80% Enterprise, без живых людей», но нам нужна конкретика.',
    [
        'Какие именно направления включены? (юрист, психолог, финансы, медицина — что ещё?)',
        'Какие боты есть? (робот-юрист, психолог-бот — какие ещё?)',
        'Какие гайды / материалы доступны сотрудникам?',
        'Есть ли лимит обращений?',
        'Можно ли получить скриншоты приложения? (для one-pager\'а)',
    ],
    'ASAP — без этого не можем начать аутрич'
)

add_question(
    2, 'БЛОКЕР — ASAP', RED,
    'Есть ли КП или презентация для облегчённой версии?',
    'После звонка или email — ЛПР попросит «пришлите что-нибудь посмотреть». Нужен хотя бы one-pager. Мы можем сделать сами на основе позиционирования — но нужно ваше согласование.',
    [
        'Есть ли готовая презентация / КП для облегчённой версии? Если да — пришлите',
        'Если нет — мы сделаем one-pager (два варианта: для PR/GR и для HR). Сможете согласовать за 1 день?',
    ],
    'ASAP — без этого нечего отправить после первого контакта'
)

add_question(
    3, 'ВАЖНО', ORANGE,
    'Тест-драйв или прямая продажа?',
    'На созвоне обсуждали: партнёр хочет бесплатный тест-драйв → конверсию, вы склоняетесь к прямой продаже. Нам нужно решение — это определяет весь скрипт и подход.',
    [
        'Идём с прямой продажей (наша рекомендация при чеке 150-500 тыс.)?',
        'Или предлагаем тест-драйв? Если да — сколько дней, что доступно, автоматический переход в платную?',
        'Или оба варианта — прямая продажа основной, тест-драйв как fallback при сомнениях?',
    ]
)

add_question(
    4, 'ВАЖНО', ORANGE,
    'Кто принимает звонок от заинтересованной компании?',
    'Мы находим компанию, звоним, они говорят «да, интересно» — дальше кто? Нужен человек, который проведёт следующий разговор и доведёт до сделки. Без этого лиды теряются.',
    [
        'Кто принимает тёплые лиды — вы лично, партнёр в ЕКБ, или кто-то из команды?',
        'За какое время сможете связаться с заинтересованной компанией? (идеально — в тот же день)',
    ]
)

add_question(
    5, 'ВАЖНО', ORANGE,
    'Можно ли упоминать Enterprise-клиентов?',
    'Для облегчённой версии пока нет своих кейсов. Лучший аргумент: «Платформа, на которой работают Северсталь, Газпромбанк и Лента — теперь доступна среднему бизнесу». Но нужно убедиться, что это не нарушает NDA.',
    [
        'Можно ли использовать названия Северсталь, Газпромбанк, Лента в коммуникации с потенциальными клиентами?',
        'Если нет — какие названия можно? Или только «крупнейшие предприятия РФ» без имён?',
    ]
)

add_question(
    6, 'УТОЧНЕНИЕ', GREEN,
    'Как рассчитывается цена?',
    'Знаем: от 150 тыс. (до 1000 чел.) и 500 тыс. (свыше 1000). Но на звонке ЛПР спросит точнее.',
    [
        'Есть ли формула? (например, X руб. за сотрудника × количество)',
        'Или это фиксированные пакеты? (150K / 300K / 500K)',
        'Есть ли разные тарифы (базовый / расширенный)?',
    ],
    'Не блокирует старт — можно уточнить позже. Для первых звонков достаточно «от 150 тыс./год»'
)

output = '/Users/wladyslaw/Documents/Job/FirstLeadsV2/Clients/DOBROSERVICE/01_positioning/questions-for-client.docx'
doc.save(output)
print(f'Saved: {output}')
