#!/usr/bin/env python3
"""Generate DOCX positioning report for ДОБРОСЕРВИС Digital."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Color Palette ───
BLUE = RGBColor(0x25, 0x63, 0xEB)
BLUE_DARK = RGBColor(0x1D, 0x4E, 0xD8)
GREEN = RGBColor(0x05, 0x96, 0x69)
GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)
GRAY_600 = RGBColor(0x4B, 0x55, 0x63)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_800 = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BG_BLUE_LIGHT = "DBEAFE"
BG_GREEN_LIGHT = "D1FAE5"
BG_GREEN_LIGHTER = "ECFDF5"
BG_YELLOW_LIGHT = "FEF3C7"
BG_RED_LIGHT = "FEE2E2"
BG_GRAY_LIGHT = "F3F4F6"
BG_GRAY = "E5E7EB"


# ─── Helpers ───

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, side, color="000000", sz="4", val="single"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    border = parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
    )
    existing = tcBorders.find(qn(f'w:{side}'))
    if existing is not None:
        tcBorders.remove(existing)
    tcBorders.append(border)


def add_title_block(doc, title, subtitle, version, color_hex):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color_hex)
    for side in ['top', 'bottom', 'start', 'end']:
        set_cell_border(cell, side, color=color_hex, sz="0", val="none")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)
    run2.font.color.rgb = WHITE
    run2.font.name = 'Calibri'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(40)
    run3 = p3.add_run(version)
    run3.font.size = Pt(10)
    run3.font.color.rgb = WHITE
    run3.font.name = 'Calibri'

    doc.add_paragraph()


def add_heading(doc, text, level=1, color=BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=None, space_after=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p


def add_rich_para(doc, parts, space_after=None, space_before=None):
    p = doc.add_paragraph()
    for text, fmt in parts:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        if fmt.get('bold'):
            run.font.bold = True
        if fmt.get('italic'):
            run.font.italic = True
        if fmt.get('color'):
            run.font.color.rgb = fmt['color']
        if fmt.get('size'):
            run.font.size = Pt(fmt['size'])
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.name = 'Calibri'
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.name = 'Calibri'
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
    return p


def add_styled_table(doc, headers, rows, col_widths=None, header_bg=BG_BLUE_LIGHT, highlight_rows=None, highlight_bg=BG_GREEN_LIGHTER):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        set_cell_shading(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = GRAY_700

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_callout_box(doc, title, text, border_color_hex="2563EB", bg_hex=BG_BLUE_LIGHT, title_color=BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    left_cell.width = Cm(0.3)
    set_cell_shading(left_cell, border_color_hex)
    left_cell.text = ""

    right_cell = table.cell(0, 1)
    set_cell_shading(right_cell, bg_hex)

    p = right_cell.paragraphs[0]
    run_t = p.add_run(title)
    run_t.font.bold = True
    run_t.font.color.rgb = title_color
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(11)

    p2 = right_cell.add_paragraph()
    run_c = p2.add_run(text)
    run_c.font.name = 'Calibri'
    run_c.font.size = Pt(10)

    for cell in [left_cell, right_cell]:
        for side in ['top', 'bottom', 'end', 'start']:
            set_cell_border(cell, side, val="none", sz="0")

    doc.add_paragraph()
    return table


def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = GRAY_800

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


# ════════════════════════════════════════════════════════════════
# GENERATE POSITIONING REPORT
# ════════════════════════════════════════════════════════════════

def generate():
    doc = setup_doc()

    add_title_block(
        doc,
        "ДОБРОСЕРВИС Digital",
        "Product Positioning Document\nПо методологии April Dunford «Obviously Awesome»",
        "Version 1.0 · 12 марта 2026",
        "2563EB"
    )

    # ═══════════════════════════════════════
    # ONE-PAGER
    # ═══════════════════════════════════════
    add_heading(doc, "Резюме позиционирования", level=1, color=GREEN)
    add_para(doc, "Полностью цифровая программа поддержки сотрудников для промышленного среднего бизнеса в регионах", italic=True, color=GRAY_600, size=10)

    # Key numbers table
    add_styled_table(doc,
        ["Показатель", "Значение"],
        [
            ["Цена", "150-500K руб./год"],
            ["Направления поддержки", "9"],
            ["Целевой сегмент", "500-1500 сотрудников"],
            ["Конкуренция в нише", "0 (промышленный mid-market в регионах)"],
        ],
        header_bg=BG_GREEN_LIGHT,
    )

    # Positioning Statement PR/GR
    add_heading(doc, "Positioning Statement — PR/GR", level=2, color=GRAY_700)
    add_callout_box(doc,
        "",
        "Для PR/GR-директоров промышленных предприятий Свердловской области, "
        "которым нужен инфоповод и признание от власти, ДОБРОСЕРВИС Digital — это "
        "готовая программа повышения качества труда при поддержке «Единой России», "
        "которая даёт PR-контент, письмо губернатору и участие в госпроекте "
        "без собственных разработок за 150-500 тыс./год.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # Positioning Statement HR
    add_heading(doc, "Positioning Statement — HR", level=2, color=GRAY_700)
    add_callout_box(doc,
        "",
        "Для HR-директоров промышленных предприятий (500-1500 чел.), которые теряют людей "
        "из-за выгорания при урезанном бюджете, ДОБРОСЕРВИС Digital — это готовая "
        "программа поддержки сотрудников (9 направлений, 24/7), которая подключается "
        "за неделю и стоит от 150 тыс./год — дешевле штатного психолога в 3-10 раз.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # Elevator pitches
    add_heading(doc, "Elevator Pitch — PR/GR", level=2, color=GRAY_700)
    add_para(doc, "«Есть государственный проект по повышению качества труда. Ваши конкуренты уже подключаются. "
        "ДОБРОСЕРВИС — это готовая программа: 9 направлений поддержки сотрудников, письмо губернатору, "
        "PR-инфоповод. Подключение за неделю, от 150 тысяч в год. Конференция 24 марта в Екатеринбурге — "
        "можете выступить с кейсом, если подключитесь до неё.»", italic=True, color=GRAY_700, size=11)

    add_heading(doc, "Elevator Pitch — HR", level=2, color=GRAY_700)
    add_para(doc, "«100 тысяч рабочих не хватает в Свердловской области. Текучесть растёт, бюджет урезан. "
        "ДОБРОСЕРВИС — готовая программа поддержки: юрист, психолог, финансы, медицина — 9 направлений, "
        "24/7, без живых специалистов. От 150 тысяч в год — дешевле одного штатного психолога в 10 раз. "
        "Подключение за неделю.»", italic=True, color=GRAY_700, size=11)

    add_separator(doc)

    # ═══════════════════════════════════════
    # GAP ANALYSIS
    # ═══════════════════════════════════════
    add_heading(doc, "Gap Analysis", level=1, color=BLUE)
    add_para(doc, "Что мы знаем достоверно и где остались пробелы", italic=True, color=GRAY_600, size=10)

    add_styled_table(doc,
        ["Категория", "Статус", "Комментарий"],
        [
            ["Конкурентные альтернативы", "Все подтверждены", "6 альтернатив с ценами и слабостями"],
            ["Уникальные атрибуты", "Все подтверждены", "7 жёстких + 1 мягкий дифференциатор"],
            ["Value Map", "Все подтверждены", "Но proof points — только от Enterprise"],
            ["Target Segments", "Все подтверждены", "7 гипотез → 3 живых, 3 убиты, 1 отложен"],
            ["Market Category", "Все подтверждены", "Два фрейминга для двух ЛПР"],
            ["Монетизация", "Все подтверждены", "150-500K, под тендерный порог"],
            ["Proof points облегчённой версии", "ПРобел", "Нет данных — первый тест"],
            ["Тест-драйв vs прямая продажа", "Не решено", "Требуется решение Андрея"],
        ],
        header_bg=BG_BLUE_LIGHT,
        highlight_rows=[6, 7],
        highlight_bg=BG_YELLOW_LIGHT,
    )

    add_separator(doc)

    # ═══════════════════════════════════════
    # PRODUCT DEEP-DIVE
    # ═══════════════════════════════════════
    add_heading(doc, "Product Deep-Dive", level=1, color=GREEN)
    add_para(doc, "Глубокое понимание продукта — основа позиционирования", italic=True, color=GRAY_600, size=10)

    # A. Transformation Map
    add_heading(doc, "A. Transformation Map", level=2, color=GRAY_700)
    add_callout_box(doc,
        "ДО:",
        "Сотрудники промышленных предприятий сталкиваются с жизненными проблемами — юридическими, "
        "финансовыми, психологическими — и решают их сами. Или не решают. Результат: стресс → выгорание "
        "→ ошибки на производстве → текучесть. HR знает, но бюджета на EAP нет. PR/GR не видит в этом "
        "своей задачи.",
        "DC2626", BG_RED_LIGHT, RED
    )
    add_callout_box(doc,
        "ПОСЛЕ:",
        "У каждого сотрудника (и его семьи) — приложение с 9 направлениями поддержки: юрист, психолог, "
        "финансовый консультант, медицина, карьера, ЖКХ. Робот отвечает 24/7, база знаний решает 80% "
        "типовых ситуаций. Компания получает ESG-отчётность + PR-инфоповод + данные для HR.",
        "059669", BG_GREEN_LIGHT, GREEN
    )

    # B. The Hard Part
    add_heading(doc, "B. Что сложно скопировать", level=2, color=GRAY_700)
    add_callout_box(doc,
        "Конкурентный барьер:",
        "1. 19-летняя база знаний ЕЮС — миллионы обработанных кейсов. Понимаю с ~2018, Alter/Ясно — без базы\n"
        "2. White-label монополия — страховые компании уже на движке ЕЮС\n"
        "3. GR-связка — госпроект «Единая Россия», письма губернаторам\n"
        "4. Enterprise-кейсы — Северсталь, Газпромбанк, Лента (промышленные гиганты РФ)\n"
        "5. Ценовой барьер — конкуренты не могут снизить цену до 500K (у них живые специалисты)",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # C. Triggers
    add_heading(doc, "C. Триггеры покупки", level=2, color=GRAY_700)
    for item in [
        "Госпроект «Единая Россия» по повышению качества труда → PR/GR ищет, что подключить",
        "Конференция 24 марта в ЕКБ → дедлайн для подключения",
        "Рост текучести/травматизма → HR ищет инструмент удержания",
        "73% предприятий Свердловской области внедряют ESG → нужен конкретный пункт в отчётности",
        "Компания слышала, что конкуренты подключили → fear of missing out",
    ]:
        add_bullet(doc, item)

    # D. Hidden Value
    add_heading(doc, "D. Скрытая ценность", level=2, color=GRAY_700)
    for title, desc in [
        ("Входной продукт → полная версия: ", "облегчённая версия за 150-500K — воронка в Enterprise-подписку (5-7 млн)"),
        ("GR-обвязка: ", "не просто продукт, а участие в госпроекте — PR, письмо губернатору, признание"),
        ("Данные для HR: ", "аналитика обращений = обоснование перед руководством"),
        ("Семья: ", "подключаются все сотрудники + родственники — усиление perceived value"),
    ]:
        add_bullet(doc, desc, bold_prefix=title)

    add_separator(doc)

    # ═══════════════════════════════════════
    # 1. COMPETITIVE ALTERNATIVES
    # ═══════════════════════════════════════
    add_heading(doc, "1. Competitive Alternatives", level=1, color=BLUE)
    add_para(doc, "«Что клиенты делали бы, если бы вас не существовало?»", italic=True, color=GRAY_600, size=10)

    add_styled_table(doc,
        ["Альтернатива", "Сильные стороны", "Слабости / почему ДОБРОСЕРВИС лучше"],
        [
            ["1. Понимаю\n(EAP-платформа, лидер РФ)",
             "Известный бренд, IT-HR, Москва",
             "600K-1M/год, 6 направлений, не работает с промышленностью в регионах"],
            ["2. Corporate Health\n(Enterprise EAP)",
             "Международная экспертиза",
             "Не работает со средним бизнесом. Цена от 1M+. Москва/СПб"],
            ["3. Alter / Ясно\n(B2C-маркетплейсы)",
             "Удобный UX, бренд",
             "Только психология (1 направление). Per capita: ~2400-2520 руб. = 1.2M при 20% утилизации"],
            ["4. Штатный психолог\n(найм в штат)",
             "Живой контакт, знание контекста",
             "480K-1.44M/год, 1 направление, нет анонимности, нет аналитики, нет масштабирования"],
            ["5. ДМС с психологом\n(расширенная страховка)",
             "Привычно, понятно",
             "30-50K/сотрудника (в 100x дороже per capita). Психология — 1% обращений. Не заменяет EAP"],
            ["6. «Ничего не делать»\n(главный конкурент)",
             "0 руб. явных затрат",
             "Потери: 63 часа продуктивности/мес. на сотрудника с проблемой. Текучесть, выгорание, травматизм"],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_heading(doc, "Ключевые выводы", level=2, color=GREEN)
    add_numbered(doc, "Ни один конкурент не занимает промышленный mid-market в регионах — все целят в Enterprise/IT в Москве", bold_prefix="")
    add_numbered(doc, "Ни один конкурент не предлагает полностью автоматизированный EAP — все через живых специалистов", bold_prefix="")
    add_numbered(doc, "ДОБРОСЕРВИС в 4-7x дешевле ближайшего конкурента при более широком покрытии (9 vs 6 направлений)", bold_prefix="")
    add_numbered(doc, "Цена под тендерный порог — уникальное преимущество, убирает барьер закупки", bold_prefix="")

    add_separator(doc)

    # ═══════════════════════════════════════
    # 2. UNIQUE ATTRIBUTES
    # ═══════════════════════════════════════
    add_heading(doc, "2. Unique Attributes", level=1, color=BLUE)
    add_para(doc, "«Что есть у вас, чего нет у альтернатив?»", italic=True, color=GRAY_600, size=10)

    add_styled_table(doc,
        ["Атрибут", "Описание", "Уникален относительно", "Тип"],
        [
            ["Полная автоматизация", "Без живых специалистов — боты + база знаний + гайды",
             "Все конкуренты", "Жёсткий"],
            ["9 направлений", "Юрист, психолог, финансы, медицина, карьера, ЖКХ и др.",
             "Alter/Ясно (1), Corporate Health (3-4), Понимаю (6+)", "Жёсткий"],
            ["Цена 150-500K", "Под тендерный порог — решение за дни, не месяцы",
             "Все EAP-конкуренты (600K-1M+)", "Жёсткий"],
            ["19-летняя база знаний ЕЮС", "Миллионы обработанных кейсов",
             "Все (Понимаю с ~2018, Alter/Ясно — без базы)", "Жёсткий"],
            ["White-label инфраструктура", "Страховые компании уже на движке ЕЮС",
             "Все конкуренты", "Жёсткий"],
            ["GR/PR-обвязка", "Госпроект ЕР + письмо губернатору + признание от власти",
             "Все конкуренты (ни один не имеет GR-связки)", "Жёсткий"],
            ["Enterprise-кейсы", "Северсталь, Газпромбанк, Лента",
             "Alter/Ясно (IT), Corporate Health (международные)", "Жёсткий"],
            ["Per capita 300-1000 руб./год", "В 4-7x дешевле ближайшего конкурента",
             "Понимаю (1200-2000), Alter (~2400), Ясно (~2520)", "Мягкий"],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_heading(doc, "Strategic Differentiation Narrative", level=2, color=GRAY_700)

    add_callout_box(doc,
        "1. Почему это невозможно скопировать?",
        "• 19 лет данных ЕЮС — Понимаю с ~2018, Alter/Ясно — B2C-маркетплейсы без собственной базы. "
        "Миллионы обработанных кейсов не воспроизвести\n"
        "• White-label монополия — страховые уже на движке ЕЮС. Новому конкуренту нужно убедить страховщиков переключиться\n"
        "• GR-связка — Единая Россия, губернаторы, региональные пилоты. Стартап не воспроизведёт",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_callout_box(doc,
        "2. Почему клиент не вернётся к альтернативам?",
        "• Данные: HR-аналитика накапливает историю — переключение = потеря данных\n"
        "• GR: если начал участвовать в госпроекте, выходить неудобно (публичное обязательство)\n"
        "• Enterprise: если перешёл в полную версию (5-7 млн) — зависимость от живых специалистов ЕЮС",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_callout_box(doc,
        "3. Moat Analysis",
        "• Data accumulation: 19 лет — невоспроизводимый массив знаний\n"
        "• Platform lock-in: страховые компании на движке ЕЮС\n"
        "• Government partnership: GR-связка с партией и губернаторами\n"
        "• Price barrier: конкуренты не могут снизить цену до 500K — у них живые специалисты",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_separator(doc)

    # ═══════════════════════════════════════
    # 3. VALUE MAP
    # ═══════════════════════════════════════
    add_heading(doc, "3. Value Map", level=1, color=BLUE)
    add_para(doc, "«Какую ценность эти уникальные атрибуты создают для клиентов?»", italic=True, color=GRAY_600, size=10)

    add_styled_table(doc,
        ["Атрибут", "Что даёт", "Почему важно", "Уверенность"],
        [
            ["База знаний ЕЮС", "Бот/гайды решают 80% типовых ситуаций 24/7",
             "Мгновенный ответ сотруднику без живого человека", "Claimed"],
            ["Цена под тендер", "Решение принимается за дни",
             "Убирает главный барьер — закупочную процедуру. До 500 тыс. = решение директора", "Proven"],
            ["GR/PR-обвязка", "Готовый инфоповод, коммуникация с властью, признание",
             "Для PR/GR — «нулевой продукт», не нужно ничего создавать с нуля", "Proven"],
            ["Enterprise-кейсы", "Северсталь, Газпромбанк, Лента используют",
             "Снижает risk perception: «если они пользуются — мы тоже можем»", "Proven"],
            ["Готовая соцпрограмма", "Всё в одном приложении, подключение за дни",
             "Единственный способ иметь well-being для среднего бизнеса без бюджета", "Claimed"],
            ["HR-аналитика", "Данные для обоснования перед руководством",
             "Без данных — «потратили деньги непонятно на что»", "Inferred"],
            ["Per capita «как корп. питание»", "3500 руб./сотрудник/год — в 10x дешевле штатного психолога",
             "«Это не инвестиция, это расходник» — не требует серьёзного обоснования", "Proven"],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_callout_box(doc,
        "Ключевая проблема Value Map",
        "Все измеримые результаты (снижение текучести, больничных, выгорания) — от Enterprise-версии "
        "с живыми специалистами. Для облегчённой версии результатов пока НЕТ. Это первый тест. "
        "Нужно собрать proof points по итогам первых подключений.",
        "DC2626", BG_RED_LIGHT, RED
    )

    add_separator(doc)

    # ═══════════════════════════════════════
    # 4. TARGET SEGMENTS
    # ═══════════════════════════════════════
    add_heading(doc, "4. Target Segments", level=1, color=BLUE)
    add_para(doc, "«Кому эта уникальная ценность важнее всего?»", italic=True, color=GRAY_600, size=10)
    add_rich_para(doc, [
        ("Исследовано ", {}),
        ("7 сегментных гипотез", {"bold": True}),
        (". 3 сегмента живы, 3 убиты с доказательствами, 1 отложен.", {}),
    ])
    add_callout_box(doc,
        "Макросигнал:",
        "Дефицит 100 тыс. рабочих в Свердловской области. Травматизм +14.7% г/г. "
        "73% предприятий Свердловской области внедряют ESG-практики. 59% компаний расширяют well-being. "
        "Конференция 24.03 в ЕКБ — первая точка продаж.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # Segments priority table
    add_heading(doc, "Приоритеты сегментов", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Сегмент", "Приоритет", "Что подтверждено", "Что нужно проверить"],
        [
            ["Промышленность Свердл. обл. → PR/GR", "Приоритет 1",
             "40-60 компаний, 0 конкурентов, GR-повод, СОСПП (500+)", "PR/GR быстрее принимает решение"],
            ["Промышленность Свердл. обл. → HR", "Приоритет 1",
             "Дефицит 100K рабочих, +14.7% травматизм, 59% well-being", "HR готов тратить при урезанном бюджете"],
            ["Строительство Свердл. обл. → ГД", "Приоритет 2",
             "Травматизм, охрана труда", "ГД готов платить за соцпрограмму"],
        ],
        header_bg=BG_BLUE_LIGHT,
        highlight_rows=[0, 1],
        highlight_bg=BG_GREEN_LIGHTER,
    )

    # P1: PR/GR segment card
    add_heading(doc, "Приоритет 1 (A/B-тест): PR/GR-директор", level=2, color=GREEN)
    add_para(doc, "Промышленные предприятия Свердловской области → PR/GR-директор", bold=True)

    add_rich_para(doc, [
        ("JTBD: ", {"bold": True, "color": GRAY_600}),
        ("«Когда появляется госинициатива по повышению качества труда, я хочу быстро подключить "
         "готовое решение, чтобы получить признание от губернатора и PR-инфоповод без собственных разработок.»", {"italic": True}),
    ])

    add_rich_para(doc, [
        ("Профиль: ", {"bold": True, "color": GRAY_600}),
        ("Металлургия, машиностроение, трубное производство, приборостроение. 500-1500 чел. "
         "Свердловская область. Члены СОСПП (500+ организаций) или Уральской ТПП (900+ членов).", {}),
    ])

    add_heading(doc, "Почему покупают:", level=3, color=GRAY_700)
    add_numbered(doc,
        " Госпроект «Единая Россия» по повышению качества труда. Письмо губернатору, "
        "признание на уровне области. PR/GR-директор получает готовый инфоповод без собственных разработок. "
        "→ ДОБРОСЕРВИС = готовая обвязка: продукт + коммуникация с властью + PR-контент.",
        bold_prefix="Есть государственный проект — нужно участвовать.")
    add_numbered(doc,
        " Событие в Екатеринбурге, промышленные компании региона. Кто подключился до "
        "конференции — может выступить с кейсом. → Триггер: «подключитесь сейчас — выступите на конференции».",
        bold_prefix="Конференция 24 марта — нужно показать результат.")
    add_numbered(doc,
        " Программа поддержки сотрудников = часть ESG-стратегии. "
        "Без ДОБРОСЕРВИС — нет инструмента, есть только декларация. "
        "→ 150-500 тыс. = конкретный пункт в ESG-отчётности.",
        bold_prefix="73% предприятий Свердловской области уже внедряют ESG.")

    add_rich_para(doc, [
        ("Конкурентный ландшафт: ", {"bold": True, "color": GRAY_600}),
        ("Нулевая конкуренция.", {"bold": True}),
        (" Ни один EAP-провайдер не работает с промышленным mid-market в Свердловской области. "
         "Понимаю — Москва/IT. Corporate Health — Enterprise. Alter/Ясно — B2C-маркетплейсы.", {}),
    ])

    add_rich_para(doc, [
        ("Размер сегмента: ", {"bold": True, "color": GRAY_600}),
        ("40-60 предприятий нужного размера (500-1500 чел.) в Свердловской области. 90+ если до 2000 чел.", {}),
    ])

    doc.add_paragraph()

    # P1 (B): HR segment card
    add_heading(doc, "Приоритет 1 (A/B-тест): HR-директор", level=2, color=GREEN)
    add_para(doc, "Промышленные предприятия Свердловской области → HR-директор", bold=True)

    add_rich_para(doc, [
        ("JTBD: ", {"bold": True, "color": GRAY_600}),
        ("«Когда текучесть растёт, а бюджет урезан, я хочу внедрить готовую программу поддержки, "
         "чтобы снизить увольнения без собственной разработки.»", {"italic": True}),
    ])

    add_heading(doc, "Почему покупают:", level=3, color=GRAY_700)
    add_numbered(doc,
        " Текучесть растёт, найм дорожает. HR нужен инструмент удержания, но бюджет урезан. "
        "→ ДОБРОСЕРВИС = готовая соцпрограмма за «погрешность в налогах».",
        bold_prefix="Дефицит 100 тыс. рабочих в Свердловской области.")
    add_numbered(doc,
        " Выгорание и стресс — 2-я причина увольнений (34% по HH.ru). HR обязан что-то делать, "
        "но разработать своё — нереально. → Готовое решение: 9 направлений, подключение за дни.",
        bold_prefix="Травматизм +14.7% г/г.")
    add_numbered(doc,
        " Тренд на благополучие сотрудников — HR должен показать руководству результат. "
        "→ ДОБРОСЕРВИС + аналитика = конкретный KPI для HR.",
        bold_prefix="59% компаний расширяют well-being.")

    add_rich_para(doc, [
        ("Почему это Приоритет 1, но вторая гипотеза: ", {"bold": True, "color": GRAY_600}),
        ("HR медленнее принимает решения, бюджет урезан, нет той же urgency что у PR/GR. "
         "Тестируем параллельно на тех же компаниях — данные покажут.", {}),
    ])

    doc.add_paragraph()

    # P2: Construction
    add_heading(doc, "Приоритет 2: Строительство", level=2, color=GRAY_700)
    add_para(doc, "Строительные компании Свердловской обл. → Генеральный директор", bold=True)

    add_rich_para(doc, [
        ("Профиль: ", {"bold": True, "color": GRAY_600}),
        ("Строительные компании 200-1000 чел. Свердловская область. "
         "Высокий травматизм, низкая квалификация рабочих.", {}),
    ])
    add_rich_para(doc, [
        ("Почему Приоритет 2: ", {"bold": True, "color": GRAY_600}),
        ("Более низкая оценка (30/50). Нет HR-структуры для внедрения. "
         "ГД занят, порог входа выше. Отложен до получения первых результатов по P1.", {}),
    ])

    # Strategy box
    add_callout_box(doc,
        "Стратегия выхода",
        "Приоритет 1 (A/B-тест): PR/GR-директор + HR-директор → те же самые компании\n"
        "  → Приоритет 2 (после P1): Строительство → через ГД\n\n"
        "Те же предприятия, два ЛПР. Цель: первые ответы, discovery-звонки, приглашения на конференцию 24.03.",
        "059669", BG_GREEN_LIGHT, GREEN
    )

    # Rejected segments
    add_heading(doc, "Отклонённые и отложенные сегменты", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Сегмент", "Статус", "Причина"],
        [
            ["Горнодобыча", "Убит",
             "Компании слишком крупные (5000+ чел.) — это рынок Enterprise-версии. FL не может пройти закупку"],
            ["Пищевая промышленность", "Убит",
             "Слабая боль well-being. Нет GR-связки. Маржинальность ниже металлургии"],
            ["Транспорт", "Убит",
             "Нет HR-структуры. Вахтовый режим — не будут пользоваться приложением"],
            ["Химическая промышленность", "Отложен",
             "Менее 20 компаний нужного размера в Свердловской области. Вернуться при расширении"],
        ],
        header_bg=BG_GRAY,
    )

    add_separator(doc)

    # ═══════════════════════════════════════
    # 5. MARKET CATEGORY
    # ═══════════════════════════════════════
    add_heading(doc, "5. Market Category", level=1, color=BLUE)
    add_para(doc, "«Какая рамка восприятия делает вашу ценность очевидной?»", italic=True, color=GRAY_600, size=10)

    add_callout_box(doc,
        "Контекст",
        "Рынок EAP НЕ сформирован для среднего бизнеса. Крупные компании знают про EAP. Средние — нет. "
        "White-label означает: то, что средние видят (если видят) — это перелицованный ДОБРОСЕРВИС через страховые.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_styled_table(doc,
        ["Стратегия", "Фрейминг", "Плюсы", "Минусы"],
        [
            ["A. Head-to-Head", "EAP-платформа",
             "Понятно для HR-профессионалов", "Средний бизнес в регионах не знает слово «EAP»"],
            ["B. Niche (PR/GR)", "Участие в госпрограмме повышения качества труда",
             "Уникально, GR-связка как дифференциатор", "Узкий фрейминг, привязан к политической конъюнктуре"],
            ["C. Niche (HR)", "Готовая программа поддержки сотрудников",
             "Понятно без терминологии, акцент на «готовое решение»", "Не уникально как фрейминг"],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_callout_box(doc,
        "Рекомендация: два разных фрейминга для двух ЛПР",
        "• Для PR/GR (P1): «Участие в государственной программе повышения качества труда — "
        "бесплатное PR-сопровождение + признание от власти»\n"
        "• Для HR (P2): «Готовая программа поддержки сотрудников — подключение за неделю, от 150 тыс./год»\n\n"
        "Два разных позиционирования для двух ЛПР. A/B-тест покажет, какой фрейминг работает.",
        "059669", BG_GREEN_LIGHT, GREEN
    )

    add_separator(doc)

    # ═══════════════════════════════════════
    # 6. RELEVANT TRENDS
    # ═══════════════════════════════════════
    add_heading(doc, "6. Relevant Trends", level=1, color=BLUE)

    for name, impact in [
        ("Дефицит кадров → фокус на удержание",
         "HH.ru: выгорание и стресс — 2-я причина увольнений (34%). 100 тыс. дефицит рабочих в Свердловской области. "
         "Компании вынуждены инвестировать в удержание"),
        ("Государственное внимание к качеству труда",
         "Госпроект «Единая Россия» по внедрению технологий повышения качества труда. "
         "Письма губернаторам. Конференция 24.03 в ЕКБ"),
        ("Сокращение бюджетов → спрос на «готовое»",
         "Разработать свою соцпрограмму — дорого. Компании ищут готовые решения. "
         "ДОБРОСЕРВИС = «из коробки» за 350 тыс."),
        ("HR-бюджет режут первым",
         "HR не может защитить новую статью расходов → нужен другой бюджетодержатель. "
         "PR/GR-директор как альтернативный ЛПР — инсайт Андрея"),
    ]:
        add_bullet(doc, impact, bold_prefix=name + ": ")

    add_separator(doc)

    # ═══════════════════════════════════════
    # MONETISATION
    # ═══════════════════════════════════════
    add_heading(doc, "Монетизация", level=1, color=BLUE)

    add_heading(doc, "Модель ценообразования", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Размер компании", "Цена/год", "Per capita", "Закупочная процедура"],
        [
            ["100-1000 сотрудников", "от 150 тыс. руб.", "150-3500 руб./чел.", "Без тендера, решение директора"],
            ["Свыше 1000 сотрудников", "500 тыс. руб. (потолок)", "50-500 руб./чел.", "Без тендера, решение директора"],
            ["Enterprise (полная версия)", "5-7 млн руб.", "—", "Тендер, цикл до года"],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_heading(doc, "Ценовой якорь", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Альтернатива", "Стоимость", "ДОБРОСЕРВИС Digital"],
        [
            ["ДОБРОСЕРВИС Enterprise", "5-7 млн/год", "В 10-40x дешевле"],
            ["Штатный психолог", "~1.2-1.8 млн/год", "В 3-10x дешевле + 8 направлений сверху"],
            ["Alter/Ясно (маркетплейс)", "~1.2M при 20% утилизации", "В 4-7x дешевле, 9 направлений vs 1"],
            ["Понимаю (EAP)", "600K-1M", "В 4-7x дешевле, без тендера"],
            ["Ничего не делать", "0 руб. (явные)", "150-500 тыс. = «погрешность в налогах»"],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_callout_box(doc,
        "Стратегическая логика",
        "1. Потолок 500 тыс. = под тендерный порог. Решение за дни, а не месяцы\n"
        "2. Облегчённая → полная версия. Входной продукт — воронка в полную подписку (5-7 млн)\n"
        "3. «Погрешность в налогах». Настолько мало, что не требует серьёзного обоснования\n"
        "4. Подключаются все + родственники. Нет лимита пользователей → усиление perceived value\n"
        "5. Тест-драйв vs. прямая продажа — не решено, тестируем оба варианта",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_separator(doc)

    # ═══════════════════════════════════════
    # ASSUMPTIONS
    # ═══════════════════════════════════════
    add_heading(doc, "Допущения (требуют валидации аутричем)", level=1, color=ORANGE)

    for item in [
        "PR/GR быстрее принимает решение, чем HR — гипотеза Андрея, ключевая для всей стратегии. Тестируем A/B",
        "«Дёшево = плохо» — цена 4-7x ниже рынка может вызвать недоверие. Слушаем возражения на звонках",
        "GR-повод как триггер покупки — госпроект реально ускоряет решение? Или «ещё одно письмо»?",
        "80% покрытие задач без живых специалистов — заявка клиента. Насколько целевая аудитория это примет?",
        "Компании 500-1500 знают, что такое EAP — если не знают, нужен другой фрейминг (именно поэтому не используем термин «EAP»)",
    ]:
        add_bullet(doc, item)

    add_separator(doc)

    # ═══════════════════════════════════════
    # OPEN QUESTIONS
    # ═══════════════════════════════════════
    add_heading(doc, "Открытые вопросы", level=1, color=ORANGE)
    add_para(doc, "Вопросы для Андрея / команды вынесены в отдельный документ: questions-for-client.docx",
             color=GRAY_700)

    add_separator(doc)

    # ═══════════════════════════════════════
    # OPERATIONAL PLAN
    # ═══════════════════════════════════════
    add_heading(doc, "Операционный план FL", level=1, color=BLUE)

    add_styled_table(doc,
        ["#", "Задача", "Срок"],
        [
            ["1", "Согласовать позиционирование (этот документ) с Андреем", "13-14.03"],
            ["2", "Подготовить 2 one-pager'а (PR/GR + HR)", "14-15.03"],
            ["3", "Собрать таргет-лист: промышленные компании Свердл. обл. 500-1500 чел.", "14-17.03"],
            ["4", "Написать 2 скрипта (PR/GR-заход + HR-заход)", "15-17.03"],
            ["5", "Начать аутрич (email + звонки)", "17-18.03"],
            ["6", "Конференция в ЕКБ (Точка кипения)", "24.03"],
        ],
        header_bg=BG_BLUE_LIGHT,
        highlight_rows=[5],
        highlight_bg=BG_YELLOW_LIGHT,
    )

    add_heading(doc, "Сценарии результата контакта", level=2, color=GRAY_700)
    add_bullet(doc, " → передаём как тёплый лид Андрею/селзам", bold_prefix="«Да, интересно»")
    add_bullet(doc, " → приглашение на конференцию 24.03", bold_prefix="«Интересно, но не уверены»")
    add_bullet(doc, " → фиксируем возражение, причину, идём дальше", bold_prefix="«Нет»")

    add_callout_box(doc,
        "Ключевой тест",
        "Кто реагирует быстрее и конвертируется лучше — HR или PR/GR?\n"
        "Это определит всю дальнейшую стратегию продаж облегчённой версии ДОБРОСЕРВИС.",
        "059669", BG_GREEN_LIGHT, GREEN
    )

    # Footer
    doc.add_paragraph()
    add_para(doc, "Product Positioning Document v1.0 | Методология April Dunford \"Obviously Awesome\" | 12 марта 2026",
             color=GRAY_600, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Следующий шаг: Offer Architect → Entry Strategist → GTM Compiler",
             color=GRAY_600, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    output = '/Users/wladyslaw/Documents/Job/FirstLeadsV2/Clients/DOBROSERVICE/01_positioning/positioning-report-v1.docx'
    doc.save(output)
    print(f'Saved: {output}')


if __name__ == '__main__':
    generate()
