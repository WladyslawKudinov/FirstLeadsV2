#!/usr/bin/env python3
"""Generate DOCX versions of Turn 4 segment artifacts for RAG Product."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Colors ───
BLUE = RGBColor(0x25, 0x63, 0xEB)
BLUE_DARK = RGBColor(0x1D, 0x4E, 0xD8)
GREEN = RGBColor(0x05, 0x96, 0x69)
GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
GRAY_600 = RGBColor(0x4B, 0x55, 0x63)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_800 = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BG_BLUE = "DBEAFE"
BG_GREEN = "D1FAE5"
BG_GREEN2 = "ECFDF5"
BG_YELLOW = "FEF3C7"
BG_RED = "FEE2E2"
BG_PURPLE = "F5F3FF"
BG_GRAY = "F3F4F6"
BG_GRAY2 = "E5E7EB"


# ─── Helpers ───

def shd(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def bdr(cell, side, color="000000", sz="4", val="single"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = parse_xml(f'<w:tcBorders {nsdecls("w")}/>'); tcPr.append(tcB)
    b = parse_xml(f'<w:{side} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
    ex = tcB.find(qn(f'w:{side}'))
    if ex is not None: tcB.remove(ex)
    tcB.append(b)

def setup():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Calibri'; s.font.size = Pt(10.5)
    s.font.color.rgb = GRAY_800
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.15
    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']; hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        if lv == 1: hs.font.size = Pt(18); hs.paragraph_format.space_before = Pt(24)
        elif lv == 2: hs.font.size = Pt(14); hs.paragraph_format.space_before = Pt(18)
        else: hs.font.size = Pt(11); hs.paragraph_format.space_before = Pt(12)
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    return doc

def title_block(doc, title, subtitle, meta, color_hex):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); shd(c, color_hex)
    for s in ['top','bottom','start','end']: bdr(c, s, color=color_hex, sz="0", val="none")
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(36)
    r = p.add_run(title); r.font.size = Pt(24); r.font.color.rgb = WHITE; r.bold = True; r.font.name = 'Calibri'
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle); r2.font.size = Pt(13); r2.font.color.rgb = WHITE; r2.font.name = 'Calibri'
    p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; p3.space_after = Pt(36)
    r3 = p3.add_run(meta); r3.font.size = Pt(10); r3.font.color.rgb = WHITE; r3.font.name = 'Calibri'
    doc.add_paragraph()

def h(doc, text, level=1, color=None):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        if color: r.font.color.rgb = color
        r.font.name = 'Calibri'
    return hd

def p(doc, text, bold=False, italic=False, color=None, size=None):
    pa = doc.add_paragraph(); r = pa.add_run(text); r.font.name = 'Calibri'
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return pa

def rich(doc, parts, space_after=None):
    """parts = [(text, {bold, italic, color, size}), ...]"""
    pa = doc.add_paragraph()
    for text, fmt in parts:
        r = pa.add_run(text); r.font.name = 'Calibri'
        if fmt.get('b'): r.bold = True
        if fmt.get('i'): r.italic = True
        if fmt.get('c'): r.font.color.rgb = fmt['c']
        if fmt.get('s'): r.font.size = Pt(fmt['s'])
    if space_after is not None: pa.paragraph_format.space_after = Pt(space_after)
    return pa

def bullet(doc, text, bold_prefix=None):
    pa = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = pa.add_run(bold_prefix); r.bold = True; r.font.name = 'Calibri'
        pa.add_run(text).font.name = 'Calibri'
    else:
        pa.clear(); pa.add_run(text).font.name = 'Calibri'
    return pa

def numbered(doc, text, bold_prefix=None):
    pa = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = pa.add_run(bold_prefix); r.bold = True; r.font.name = 'Calibri'
        pa.add_run(text).font.name = 'Calibri'
    else:
        pa.clear(); pa.add_run(text).font.name = 'Calibri'
    return pa

def tbl(doc, headers, rows, col_widths=None, header_bg=BG_BLUE):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hd; shd(c, header_bg)
        for pp in c.paragraphs:
            for r in pp.runs: r.bold = True; r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = GRAY_700
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size = Pt(9); r.font.name = 'Calibri'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def callout(doc, title, text, border_hex="2563EB", bg_hex=BG_BLUE, title_color=BLUE):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = t.cell(0, 0); lc.width = Cm(0.3); shd(lc, border_hex); lc.text = ""
    rc = t.cell(0, 1); shd(rc, bg_hex)
    pp = rc.paragraphs[0]; r = pp.add_run(title); r.bold = True; r.font.color.rgb = title_color; r.font.name = 'Calibri'; r.font.size = Pt(10)
    p2 = rc.add_paragraph(); r2 = p2.add_run(text); r2.font.name = 'Calibri'; r2.font.size = Pt(9.5)
    for c in [lc, rc]:
        for s in ['top','bottom','end','start']: bdr(c, s, val="none", sz="0")
    doc.add_paragraph()

def sep(doc):
    pp = doc.add_paragraph()
    pPr = pp._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/></w:pBdr>'))

def reason(doc, num, headline, evidence, outcome):
    """Add a 'why they buy' reason."""
    rich(doc, [
        (f"{num}. ", {'b': True, 'c': RED, 's': 11}),
        (headline, {'b': True, 's': 10.5}),
    ])
    p(doc, evidence, italic=True, color=GRAY_600, size=9.5)
    rich(doc, [("→ ", {'b': True, 'c': GREEN}), (outcome, {'c': GREEN_DARK, 's': 9.5})], space_after=8)


# ════════════════════════════════════════════════════════════════
# 02_Segments_Reachable.docx
# ════════════════════════════════════════════════════════════════

def segment_block(doc, letter, name, jtbd, profile_rows, lpr_rows, channels_rows,
                  list_steps, verdict, examples_headers, examples_rows, extra_note=None):
    """Reusable segment block for 02_Segments_Reachable."""
    h(doc, f"Segment {letter}: {name}", 1, BLUE)
    callout(doc, "JTBD", jtbd, "2563EB", BG_BLUE, BLUE)

    h(doc, "Профиль компании (observable, filterable)", 2)
    tbl(doc, ["Параметр", "Значение", "Источник"], profile_rows, col_widths=[3, 10, 4])

    h(doc, "ЛПР", 2)
    tbl(doc, ["Параметр", "Значение"], lpr_rows, col_widths=[4, 13])

    h(doc, "FL Outreach Channels", 2)
    tbl(doc, ["Канал", "Что FL делает", "Результат"], channels_rows, col_widths=[3, 11, 3])

    h(doc, "Как построить список 50 компаний за 30 минут", 3)
    for i, step in enumerate(list_steps):
        numbered(doc, f" {step}", bold_prefix=f"{i+1}.")
    rich(doc, [("Вердикт: ", {'b': True}), (verdict, {'c': GREEN_DARK})], space_after=8)

    h(doc, "Компании-примеры (топ-5)", 3)
    tbl(doc, examples_headers, examples_rows, header_bg=BG_GREEN)

    if extra_note:
        p(doc, extra_note, italic=True, color=GRAY_500, size=9.5)

    sep(doc)


def generate_02():
    doc = setup()
    title_block(doc,
        "Сегменты → Как достать",
        "02_Segments_Reachable • RAG Product",
        "Март 2026 • REDO v2",
        "2563EB")

    # ── IT ──
    segment_block(doc,
        "A", "IT-компании / Производители ПО",
        "Документация устаревает, онбординг долгий, support перегружен, ChatGPT заблокирован в РФ → нужен AI-ассистент для внутренних знаний с on-prem.",
        profile_rows=[
            ["Отрасль", "Разработка ПО, SaaS, IT-аутсорс (ОКВЭД 62.01, 62.02, 63.11.1)", "[E] Rusprofile"],
            ["Размер", "20-150 сотрудников, выручка 50-300M₽", "[E] HighTime SaaS rating"],
            ["Техно-стек", "Собственный продукт/платформа, внутренняя документация, support-команда", "[E] HH.ru вакансии"],
            ["Регион", "Москва, СПб (основная масса), Екатеринбург, Новосибирск", "[E] firmoteka.ru"],
            ["Адресуемый рынок", "~94 SaaS-компании в рейтинге + ~200-300 по ОКВЭД 62.01", "[E] HighTime + Rusprofile"],
        ],
        lpr_rows=[
            ["Должность", "CTO / VP Engineering / Технический директор (50+ чел.). CEO (<50 чел.)"],
            ["KPI", "Скорость разработки, time-to-market, удержание команды, качество support"],
            ["Где активен", "Telegram (primary). LinkedIn заблокирован в РФ с 2016"],
            ["Конкретные персоны", "Самат Галимов (@ctodaily), Илья Чекальский (@techdir), CTO On Live (@ctoonlive)"],
            ["Email-паттерны", "firstname@company.ru, firstname.lastname@company.ru"],
        ],
        channels_rows=[
            ["Email (cold)", "firmoteka.ru/top-saas-hightime → email CTO. Hook: «Растёте на 30-50% — как с онбордингом?»", "5-10% reply"],
            ["Telegram DM", "CTO в @ctodaily, @techdir — личное сообщение с hook по компании", "3-8% reply"],
            ["HH.ru", "«technical writer» / «knowledge manager» → компания = наш таргет", "Персонализация"],
            ["Телефон", "Ресепшн → CTO (плоская структура IT-компаний)", "Демо"],
        ],
        list_steps=[
            "firmoteka.ru/top-saas-hightime → фильтр 50-300M₽, 20-150 чел. → ~25-30 (10 мин)",
            "Rusprofile ОКВЭД 62.01 + выручка 50-300M₽ → ~20+ дополнительных (10 мин)",
            "HH.ru «technical writer» + «knowledge manager» → кросс-проверка на pain signal (10 мин)",
        ],
        verdict="YES — список из 50 компаний за 15-20 минут. [E]",
        examples_headers=["#", "Компания", "Город", "Выручка", "Сотр.", "Pain signal"],
        examples_rows=[
            ["1", "Cleverence", "Москва", "296M₽", "93", "Сложная документация по мобильным платформам"],
            ["2", "Aspro", "Челябинск", "276M₽", "79", "+50% рост, растущая команда"],
            ["3", "Smartnut", "Екатеринбург", "255M₽", "47", "Борьба с обновлением документации"],
            ["4", "Textback", "Москва", "252M₽", "37", "Быстрый рост, знания не масштабируются"],
            ["5", "Sipuni", "Москва", "243M₽", "99", "+30% рост, support без базы знаний"],
        ])

    # ── Wholesale ──
    segment_block(doc,
        "B", "Оптовая торговля / Дистрибьюторы",
        "Менеджеры по продажам тратят 70% времени на рутину (КП, переписка, претензии). AI автоматизирует → больше времени на продажи.",
        profile_rows=[
            ["Отрасль", "Оптовая торговля, дистрибуция (ОКВЭД 46.x)", "[E] Rusprofile"],
            ["Размер", "30-200 сотрудников, выручка 100-500M₽", "[E] positioning v5"],
            ["Техно-стек", "CRM (Битрикс24 / AmoCRM / 1С), отдел продаж 5+ менеджеров", "[E] HH.ru"],
            ["Адресуемый рынок", "2,000-5,000 компаний по РФ, 500-1,500 Мск/СПб", "[E] Rusprofile"],
        ],
        lpr_rows=[
            ["Должность", "Коммерческий директор / РОП (50+ чел.). Собственник/CEO (<50 чел.)"],
            ["KPI", "Выручка, конверсия, средний чек, продуктивность менеджеров"],
            ["Где активен", "Telegram (primary), VK. LinkedIn слабый"],
            ["Email-паттерны", "ivan@company.ru, sales@company.ru"],
        ],
        channels_rows=[
            ["Email (cold)", "Rusprofile → email. «Вижу набираете 4 менеджера на HH — AI может закрыть рутину»", "5-8% reply"],
            ["Telegram DM", "@optlist_chat (35K), @dirclub (17K) — личное сообщение", "3-5% reply"],
            ["Телефон", "Основной номер → комм. директор. Привыкли к cold calls", "Демо"],
            ["HH.ru", "«Менеджер оптовых продаж» → 972+ вакансий → сигнал роста", "Персонализация"],
            ["CRM-интеграторы", "CRM Academy, Intervolga, Ингруппа → рекомендации клиентам", "Warm referral"],
        ],
        list_steps=[
            "Rusprofile Pro ОКВЭД 46.x + выручка 100-500M₽ + 30-200 чел. → экспорт 50 (10 мин)",
            "ExportBase (export-base.ru) — база 2,459 дистрибьюторов → фильтр в Excel (5 мин)",
            "HH.ru «менеджер по оптовым продажам» → компании-наниматели (10 мин)",
        ],
        verdict="YES — 50 компаний за 15-25 минут. [E]",
        examples_headers=["#", "Компания", "Город", "Сегмент", "Pain signal"],
        examples_rows=[
            ["1", "ГК «КСК»", "СПб", "Стройматериалы", "Нанимает менеджеров для cold sales + подготовка КП"],
            ["2", "ГК KOLOBOX", "Н. Новгород", "Шины/диски, топ-10", "110+ точек, IT-инвестиции"],
            ["3", "«ГРАТ-ВЕСТ»", "Москва", "FMCG (игрушки)", "Нанимает региональных менеджеров"],
            ["4", "«ВЕКТОР» (Total Tools)", "Москва", "Инструмент", "Развитие представительства — рост"],
            ["5", "«СМАРТ-ПРОЕКТ»", "Москва", "Стройматериалы", "Расширение, нужен опыт продаж"],
        ],
        extra_note="Плюс из v1: Русич, РБМ-СТРОЙ, МедБаза, НВ-Лаб и др. (фильтровать по профилю 100-500M₽).")

    # ── Manufacturing ──
    segment_block(doc,
        "C", "Производство с офисным блоком",
        "Инженеры 60% на нецелевых задачах. Техдокументация разбросана. Уход эксперта = потеря знаний. AI-поиск по документации + сохранение знаний.",
        profile_rows=[
            ["Отрасль", "Обрабатывающее производство (ОКВЭД 25-30)", "[E] Rusprofile"],
            ["Размер", "50-300 сотрудников, выручка 200-1000M₽", "[E] positioning v5"],
            ["Техно-стек", "КОМПАС-3D, НЕТ PLM/PDM, файловые серверы", "[E] АСКОН кейсы"],
            ["Адресуемый рынок", "3,000-5,000 компаний по РФ", "[E] Rusprofile"],
        ],
        lpr_rows=[
            ["Должность", "Технический директор / Главный инженер (рекомендатель). Генеральный директор (бюджет)"],
            ["KPI", "Сроки производства, качество, себестоимость, онбординг инженеров"],
            ["Где активен", "НИГДЕ цифрово. LinkedIn мёртв. Telegram — минимально. Канал: телефон и личные встречи"],
        ],
        channels_rows=[
            ["Телефон (PRIMARY)", "Звоним на приёмную завода (2GIS/сайт) → технический директор", "Демо"],
            ["Email (cold)", "info@ с сайта → «Видел на АСКОН что используете КОМПАС — аналогов AI в РФ нет»", "3-5% reply"],
            ["АСКОН-партнёры", "30 офисов, знают каждого клиента КОМПАС в регионе. 8-800-700-00-78", "Warm referral"],
            ["HH.ru", "«Инженер-конструктор КОМПАС-3D» → 3,146 вакансий → сигнал роста", "Персонализация"],
            ["Форумы АСКОН", "forum.ascon.ru, cccp3d.ru — компании с проблемами документации", "Lead intelligence"],
        ],
        list_steps=[
            "best.ascon.ru/gallery/ — 63+ компании-участники конкурса АСов → 20-30 (10 мин)",
            "Rusprofile ОКВЭД 25-30 + выручка 200-1000M₽ + 50-300 чел. → 20+ (10 мин)",
            "HH.ru «инженер-конструктор КОМПАС-3D» → компании-наниматели (10 мин)",
        ],
        verdict="YES — 50 компаний за 20-30 минут. [E]",
        examples_headers=["#", "Компания", "Город", "Сектор", "Pain signal"],
        examples_rows=[
            ["1", "Челябинский компрессорный завод", "Челябинск", "Компрессоры", "1-е место конкурс АСКОН, КОМПАС без PLM"],
            ["2", "Завод Двигатель", "СПб", "Двигатели", "Кейс АСКОН — КОМПАС + ВЕРТИКАЛЬ"],
            ["3", "Финист", "Екатеринбург", "Пищ. оборудование", "HH.ru: нанимает конструктора, рост объёмов"],
            ["4", "НПП Квант", "Ростов-на-Дону", "Приборостроение", "HH.ru: КОМПАС-3D, ЕСКД compliance"],
            ["5", "Гаврилов-Ямский машзавод АГАТ", "Ярославская обл.", "Машиностроение", "АСКОН проект, ограниченный IT-штат"],
        ],
        extra_note="Плюс из v1: ЗБО, Пожарные системы, КРМЗ, ИНФОТРАНС и др. (12 компаний).")

    # ── Law Firms ──
    segment_block(doc,
        "D", "Крупные юридические фирмы (топ-20 Право-300)",
        "88% юристов используют AI. Экспертный боттлнек. On-prem (адвокатская тайна). Ни один конкурент не покрывает полный набор.",
        profile_rows=[
            ["Отрасль", "Юридические фирмы (ОКВЭД 69.1), бухгалтерский аутсорс", "[E] Право-300, RAEX"],
            ["Размер", "30-150 юристов, 200-800M₽. Не топ-5 — слишком длинный цикл для FL", "[E] Право-300"],
            ["Адресуемый рынок", "~60-100 юр. фирм + ~20-30 бух. аутсорсеров = 80-130", "[E] Право-300 + RAEX"],
            ["Expansion", "→ Корп. юр. отделы 500M₽+ = 2,000-5,000 компаний", "[E] research"],
        ],
        lpr_rows=[
            ["Должность", "Управляющий партнёр / Генеральный директор / Старший партнёр"],
            ["KPI", "Utilization rate, revenue per partner, качество работы"],
            ["Где активен", "Telegram (primary), Zakon.ru, Право.ru. LinkedIn мёртв"],
        ],
        channels_rows=[
            ["Email (cold)", "Из Право-300 → email партнёра. «88% юристов уже в AI, но ни один — с вашими доками + on-prem»", "5-10% reply"],
            ["Telegram DM", "«Рульфы, Ильфы и Инхаусы» (5K), «ЗарбитражЫ» (1.5K)", "3-5% reply"],
            ["Телефон", "Звоним в фирму → управляющий партнёр. Юристы привыкли к деловым звонкам", "Демо"],
            ["Право.ru", "Мониторим Право-300, упоминаем позицию в outreach", "Персонализация"],
        ],
        list_steps=[
            "300.pravo.ru → топ-50 по выручке → 30+ подходящих юр. фирм (10 мин)",
            "RAEX рейтинг бухгалтерского аутсорса → топ-20 → 10+ подходящих (5 мин)",
            "Rusprofile ОКВЭД 69.1 + выручка 200M₽+ → 10+ дополнительных (10 мин)",
        ],
        verdict="YES — 50 компаний за 25 минут. Но TAM ограничен (~130 total). [E]",
        examples_headers=["#", "Компания", "Город", "Факты", "Тип"],
        examples_rows=[
            ["1", "КИАП", "Москва", "Top-10 по ширине практик, LegalTech-активны", "Юр."],
            ["2", "a.t.Legal", "Москва", "Группа 1 в арбитраже, банкротстве, недвижимости", "Юр."],
            ["3", "ЦПО Групп", "Москва", "13 номинаций 2025", "Юр."],
            ["4", "Хренов и Партнёры", "Москва", "30-й по юристам, 48-й по выручке", "Юр."],
            ["5", "Рустам Курмаев", "Москва", "Крупная литигация, основан 2017", "Юр."],
        ])

    # ── Medicine ──
    segment_block(doc,
        "E", "Частная медицина (клиники 300M₽+)",
        "120 звонков/день, 80% рутинных. Администраторы перегружены. On-prem = преимущество (ФЗ-152).",
        profile_rows=[
            ["Отрасль", "Частная медицина: многопрофильные клиники, сети (ОКВЭД 86.x)", "[E] Vademecum"],
            ["Размер", "50-300+ сотрудников, выручка 300M₽-3B₽. Не крупные сети", "[E] Vademecum TOP200"],
            ["Адресуемый рынок", "~150-200 клиник (Vademecum TOP200, порог входа 307M₽)", "[E] Vademecum"],
        ],
        lpr_rows=[
            ["Должность", "Исполнительный директор / Управляющий клиникой (primary). Директор/Собственник (sign-off)"],
            ["KPI", "Загрузка клиники, NPS, текучка персонала, выручка на врача"],
            ["Где активен", "Крайне ограниченно. Закрытые группы Telegram. Школа Мед. Бизнеса. Vademecum"],
        ],
        channels_rows=[
            ["Email (cold)", "Vademecum TOP200 → email с сайта клиники. «429 вакансий на HH — AI закроет 80% рутинных звонков»", "3-5% reply"],
            ["Телефон", "Звоним в клинику → исполнительный / управляющий директор", "Демо"],
            ["HH.ru", "«Администратор клиники» → 1,222+ вакансий в Москве. Массовый найм = боль", "Персонализация"],
        ],
        list_steps=[
            "Vademecum TOP200 (PDF, платный) → 200 клиник с выручкой 307M₽+ (5 мин)",
            "HH.ru «администратор клиники» → компании с высоким оборотом (15 мин)",
            "Rusprofile ОКВЭД 86.x + выручка 300M₽+ (10 мин)",
        ],
        verdict="YES — 50 клиник за 15-30 минут с Vademecum PDF. Без PDF — PARTIALLY (~30-40). [E]",
        examples_headers=["#", "Клиника", "Город", "Выручка", "Pain signal"],
        examples_rows=[
            ["1", "Здоровье 365", "Екатеринбург", "#38 Vademecum", "+38% рост, активный найм"],
            ["2", "СМТ-Клиника", "Екатеринбург", "#135 Vademecum", "+20% рост"],
            ["3", "Парацельс", "Екатеринбург", "#140 Vademecum", "+18% рост"],
            ["4", "Клиника Пасман", "Новосибирск", "Новичок TOP200", "Строит процессы с нуля"],
            ["5", "Авеню", "Ростов-на-Дону", "Франшиза", "Стандартизация при расширении"],
        ])

    # Reachability Warning for Medicine
    callout(doc, "Reachability Warning — Медицина",
        "Самый сложный сегмент для FL: ЛПР не на LinkedIn/Telegram, конкуренты (Chatme.ai) уже с кейсами, консервативная культура. Рекомендация: тестировать ПОСЛЕ кейсов из других сегментов.",
        "DC2626", BG_RED, RED)

    out = os.path.join(DIR, "02_Segments_Reachable.docx")
    doc.save(out)
    print(f"  ✓ {out}")


# ════════════════════════════════════════════════════════════════
# 03_Segments_Scored.docx
# ════════════════════════════════════════════════════════════════

def scored_segment(doc, letter, name, score, confidence, blockers, priorities, rows):
    """Reusable scored segment block for 03_Segments_Scored."""
    h(doc, f"Segment {letter}: {name} — Score: {score}", 1, BLUE)
    tbl(doc,
        ["#", "Criterion", "Score", "Type", "Evidence / Reasoning"],
        rows,
        col_widths=[0.8, 3, 1, 1.2, 11],
        header_bg=BG_BLUE)
    rich(doc, [
        ("Score: ", {'b': True}),
        (f"{score} ", {'b': True, 'c': BLUE}),
        (f"({confidence})", {'c': GRAY_600}),
    ])
    rich(doc, [("Blockers: ", {'b': True}), (blockers, {'c': RED if "C" in blockers else GRAY_600})])
    rich(doc, [("Top validation priorities: ", {'b': True}), (priorities, {})])
    sep(doc)


def generate_03():
    doc = setup()
    title_block(doc,
        "Оценка сегментов",
        "03_Segments_Scored • RAG Product",
        "Март 2026 • REDO v2 • Moore Beachhead Criteria",
        "7C3AED")

    # ── IT ──
    scored_segment(doc, "A", "IT-компании", "42/50",
        "E: 30/35 — high confidence | H: 12/15 — needs validation",
        "Нет.",
        "C2 (реальная боль vs «мы справляемся сами»), C3 (достаточность функционала для IT)",
        [
            ["C1", "Target customer exists", "5", "[E]", "CTO/VP Engineering с бюджетом. 94 SaaS в рейтинге HighTime 50M₽+. Cleverence (296M₽, 93 чел.), Sipuni (243M₽, 99 чел.). [firmoteka.ru, HH.ru]"],
            ["C2", "Compelling reason to buy", "4", "[H-high]", "20% времени на документацию [MTS Web Services]. 30% на поиск информации. ChatGPT заблокирован. VQ: «Сколько часов/неделю на поиск внутренней информации?»"],
            ["C3", "Whole product deliverable", "4", "[H-high]", "Полный цикл (docs + email + meetings). Minervasoft/AutoFAQ — частичные. On-prem решает безопасность. VQ: «Достаточно ли функционала?»"],
            ["C4", "No entrenched competitor", "4", "[E]", "Minervasoft (50-150K₽/мес) — KMS без AI. AutoFAQ — только support. ChatGPT Team заблокирован. [research, Habr]"],
            ["C5", "Partners & allies", "3", "[E]", "Нет партнёрского канала (нет аналога АСКОН). Можно через IT-медиа и Telegram CTO."],
            ["C6", "Distribution channel", "5", "[E]", "firmoteka.ru → 94 компании за 10 мин. CTO в Telegram. Email доступен. HH.ru pain signals. Цикл 2-6 нед. [firmoteka, HH, Telegram]"],
            ["C7", "Pricing fits budget", "5", "[E]", "80-110K₽/мес = 0.5-0.7 FTE. 0.3-1.3% выручки. ChatGPT Team недоступен → единственная альтернатива. [HighTime, salary benchmarks]"],
            ["C8", "Segment size", "4", "[E]", "~94 SaaS + ~200-300 по ОКВЭД 62.01 = 300-400 компаний. Достаточно для FL. [firmoteka, Rusprofile]"],
            ["C9", "Reference value", "4", "[H-high]", "IT-кейс → credibility для любого сегмента: «если IT купила вместо того чтобы делать сама». VQ: «Станет ли кейс аргументом для не-IT?»"],
            ["C10", "Outreach accessibility", "4", "[E]", "CTO через Telegram и email. LinkedIn мёртв, Telegram компенсирует. Список за 15-20 мин. Цикл 2-6 нед. [Telegram, firmoteka]"],
        ])

    # ── Wholesale ──
    scored_segment(doc, "B", "Оптовая торговля", "41/50",
        "E: 31/35 — high confidence | H: 10/15 — needs validation",
        "Нет.",
        "C3 (техническая валидация КП из прайса), C2 (реальная боль vs «у нас и так норм»)",
        [
            ["C1", "Target customer exists", "5", "[E]", "Комм. директор / РОП с бюджетом. 972+ вакансий «менеджер оптовых продаж» Мск. 2000-5000 компаний ОКВЭД 46.x, 100-500M₽. [HH, Rusprofile, ExportBase]"],
            ["C2", "Compelling reason to buy", "4", "[H-high]", "70% на рутине, 28-34% на продажи [43 отдела продаж]. ~400K₽/мес на 10 менеджеров. VQ: «>50% на рутину — согласны?»"],
            ["C3", "Whole product deliverable", "3", "[H-mid]", "КП из прайса заявлено, но не доказано на реальных прайсах. Нужна тех. валидация. VQ: «AI сгенерирует КП из вашего прайса за 30 сек?»"],
            ["C4", "No entrenched competitor", "4", "[E]", "BitrixGPT и amoAI НЕ генерируют КП из прайсов. SalesAI (39-56K₽/10 менеджеров) — только звонки. Gap подтверждён. [обзоры CRM, SalesAI]"],
            ["C5", "Partners & allies", "5", "[E]", "CRM Academy (300+ клиентов), Intervolga (Комус, Левенгук), Ингруппа. 5+ интеграторов. [crmacademy, intervolga, crmrating]"],
            ["C6", "Distribution channel", "5", "[E]", "Rusprofile Pro → 50 за 10 мин. ExportBase → 2459 дистрибьюторов. HH 972+. Telegram @optlist_chat 35K. [Rusprofile, ExportBase, HH]"],
            ["C7", "Pricing fits budget", "4", "[E]", "IT-бюджет 3-15M₽/год. ~1M₽/год = 6-15%. «Замена 1 менеджера» (100K₽/мес). imot.io 100-150K₽ → мы в рынке. [research]"],
            ["C8", "Segment size", "5", "[E]", "2000-5000 РФ, 500-1500 Мск/СПб. Крупнейший TAM. [Rusprofile, ExportBase]"],
            ["C9", "Reference value", "3", "[H-mid]", "Торговый кейс → другие B2B с отделами продаж, но НЕ на производство/юристов. Средняя transferability. VQ: «Какие метрики до/после убедительны?»"],
            ["C10", "Outreach accessibility", "3", "[E]", "Комм. директора слабо в цифре. Телефон + email. Работает, но <Telegram для IT. [research — no LinkedIn]"],
        ])

    # ── Manufacturing ──
    scored_segment(doc, "C", "Производство", "40/50",
        "E: 31/35 — high confidence | H: 9/15 — needs validation",
        "Нет формальных. C2 и C3 на уровне 3 — неизвестные, не блокеры.",
        "C3 (on-prem + PDF/ЕСКД), C2 (осознание проблемы техдиректорами)",
        [
            ["C1", "Target customer exists", "5", "[E]", "Технический директор с бюджетом. 3146 вакансий КОМПАС-3D. 16K клиентов АСКОН. ЗБО (895M₽, 247 чел.). [HH, ascon.ru, v1]"],
            ["C2", "Compelling reason to buy", "3", "[H-mid]", "60% на нецелевых — подтверждено, но осознание ниже. «Спроси Михалыча» — привычная модель. VQ: «Инженеры теряют время или это норма?»"],
            ["C3", "Whole product deliverable", "3", "[H-mid]", "On-prem обязателен — готовность не подтверждена. PDF-чертежи, ЕСКД — тех. валидация. VQ: «AI найдёт информацию в ваших ТУ/СТО/чертежах?»"],
            ["C4", "No entrenched competitor", "5", "[E]", "НУЛЕВАЯ конкуренция. «Нет ни одного отечественного» [Habr/LANIT, авг. 2025]. АСКОН — никакого AI. Leo.AI непригодно. [Habr, ascon]"],
            ["C5", "Partners & allies", "5", "[E]", "АСКОН: 30 офисов, 500K рабочих мест КОМПАС. Горячая линия. Softline Gold Partner. [ascon.ru/partners/]"],
            ["C6", "Distribution channel", "4", "[E]", "best.ascon.ru/gallery/ → 63+ компании. Rusprofile ОКВЭД 25-30. HH 3146 вакансий КОМПАС. 50 за 20-30 мин. [АСКОН, Rusprofile, HH]"],
            ["C7", "Pricing fits budget", "4", "[E]", "IT-бюджет 5-15M₽/год. ~1M₽/год = 7-20%. Якорь: ЛОЦМАН PLM 3-8M₽ → мы в 4-10x дешевле. [НАФИ, ЛОЦМАН pricing]"],
            ["C8", "Segment size", "5", "[E]", "3000-5000 компаний РФ. АСКОН 16K клиентов. Рынок инженерного ПО 50-55 млрд ₽. [Rusprofile, АСКОН, TAdviser]"],
            ["C9", "Reference value", "3", "[H-mid]", "Производственный кейс → другие производства, но слабо на торговлю/IT. VQ: «Кейс 'онбординг 2 мес вместо 6' убедителен?»"],
            ["C10", "Outreach accessibility", "3", "[E]", "Техдиректора НЕ в LinkedIn, минимально Telegram. Канал — телефон через приёмную. Медленнее, дороже. Компенсируется АСКОН. [research]"],
        ])

    # ── Law Firms ──
    scored_segment(doc, "D", "Крупные юр. фирмы", "39/50",
        "E: 28/35 — high confidence | H: 11/15 — needs validation",
        "C8=2 — малый TAM. Не убивает (expansion через корп. юр. отделы), но ограничивает как primary.",
        "C3 (юридическая специфика продукта), C2 (боль управляющих партнёров)",
        [
            ["C1", "Target customer exists", "5", "[E]", "Управляющий партнёр с бюджетом. 330+ фирм 100-500M₽. Таргет: ранг 20-50 Право-300 (КИАП, a.t.Legal, ЦПО Групп). [Право-300, RBC]"],
            ["C2", "Compelling reason to buy", "4", "[H-high]", "88% юристов в AI [Avito + Право.ru]. 500+ договоров = 1000-2000 чел.-часов. 86% выгорание. VQ: «Сколько часов партнёры на вопросы от ассоциатов?»"],
            ["C3", "Whole product deliverable", "3", "[H-mid]", "Юр. тексты, прецеденты, К+ — специфика. Может потребовать доработки. VQ: «AI понимает юр. терминологию?»"],
            ["C4", "No entrenched competitor", "5", "[E]", "К+ AI — только база К+. Гарант ИСКРА — только публичная. Doczilla — «недоработанный» [Habr]. Нейроюрист — 50 запросов/мес. Ни один не: внутренние доки + email + meetings + on-prem. [research, Habr]"],
            ["C5", "Partners & allies", "3", "[E]", "Нет очевидного партнёра. Заходить через Право.ru и рейтинги."],
            ["C6", "Distribution channel", "4", "[E]", "Право-300 → 50 фирм за 25 мин. Telegram «Рульфы» 5K, «ЗарбитражЫ» 1.5K. Email с сайтов. [300.pravo.ru, RAEX, Telegram]"],
            ["C7", "Pricing fits budget", "5", "[E]", "Топ-20: IT-бюджет 15-250M₽. ~1M₽ = 0.4-7%. Doczilla 87K₽/мес на 30 чел. → мы сопоставимы, шире. [Право-300, Doczilla pricing]"],
            ["C8", "Segment size", "2", "[E]", "TAM ~80-130. При 20% конверсии → 16-26 клиентов → ARR 14-24M₽. МАЛЫЙ. Expansion: корп. юр. отделы 2000-5000. [Право-300, RAEX, RBC]"],
            ["C9", "Reference value", "4", "[H-high]", "Юр. кейс → корп. юр. отделы (10x TAM). «Если АЛРУД использует — стоящий инструмент.» VQ: «Юр. reference убеждает корп. юристов?»"],
            ["C10", "Outreach accessibility", "4", "[E]", "Партнёры видимы через Право-300. Email доступен. Telegram активнее чем для торговли/производства. Цикл 1-3 мес. [Право-300, research]"],
        ])

    # ── Medicine ──
    scored_segment(doc, "E", "Частная медицина (300M₽+)", "34/50",
        "E: 21/35 — moderate confidence | H: 10/15 — needs validation",
        "C5=2 — нет партнёров/каналов. Не hard kill, но серьёзный ограничитель.",
        "C3 (интеграция с МИС), C7 (готовность платить 80-110K при альтернативах от 950₽)",
        [
            ["C1", "Target customer exists", "4", "[E]", "Исполнительный директор / управляющий. Vademecum TOP200. Таргет: позиции 30-150 (300M₽-3B₽). Здоровье 365 (#38), СМТ-Клиника (#135). [Vademecum, HH]"],
            ["C2", "Compelling reason to buy", "4", "[H-high]", "120 звонков/день, 80% рутинных. 1222+ вакансий администраторов Мск. V-AI Labs: 72% без человека. VQ: «Замените 1-2 администраторов на AI?»"],
            ["C3", "Whole product deliverable", "3", "[H-mid]", "Интеграция с МИС (МЕДИАЛОГ, Инфоклиника) — неизвестно. Специфика медицины. VQ: «Продукт интегрируется с вашей МИС?»"],
            ["C4", "No entrenched competitor", "3", "[E]", "Chatme.ai УЖЕ с Клиникой Фомина (70% записей через бот), ЕМЦ. TWIN от 950₽/мес. Конкуренция ЕСТЬ. [chatme.ai]"],
            ["C5", "Partners & allies", "2", "[E]", "Нет партнёрского канала. МИС-вендоры — потенциальные, но контакт не установлен. [research]"],
            ["C6", "Distribution channel", "3", "[E]", "Vademecum TOP200 → список за 15 мин. ЛПР не в открытых каналах. Только email + телефон. [research]"],
            ["C7", "Pricing fits budget", "3", "[E]", "IT-бюджет 3-9M₽. ~1M₽ = 11-33% — на грани. Chatme.ai значительно дешевле. TWIN от 35K₽. [Vademecum, competitor pricing]"],
            ["C8", "Segment size", "3", "[E]", "TAM ~150-200 клиник (Vademecum TOP200). Маленький, но > юристов. [Vademecum]"],
            ["C9", "Reference value", "3", "[H-mid]", "Мед. кейс слабо переносится (другой ЛПР, другой контекст). VQ: «Мед. кейс помогает продавать в другие отрасли?»"],
            ["C10", "Outreach accessibility", "3", "[E]", "ЛПР не в цифровых каналах. Email + телефон единственные. Консервативная культура. [research]"],
        ])

    # ── Summary Table ──
    h(doc, "Score Summary", 1, GREEN_DARK)
    tbl(doc,
        ["Segment", "C1[E]", "C2[H]", "C3[H]", "C4[E]", "C5[E]", "C6[E]", "C7[E]", "C8[E]", "C9[H]", "C10[E]", "Total", "Confidence"],
        [
            ["A: IT", "5", "4", "4", "4", "3", "5", "5", "4", "4", "4", "42/50", "E:30/35 H:12/15"],
            ["B: Торговля", "5", "4", "3", "4", "5", "5", "4", "5", "3", "3", "41/50", "E:31/35 H:10/15"],
            ["C: Производство", "5", "3", "3", "5", "5", "4", "4", "5", "3", "3", "40/50", "E:31/35 H:9/15"],
            ["D: Юр. фирмы", "5", "4", "3", "5", "3", "4", "5", "2", "4", "4", "39/50", "E:28/35 H:11/15"],
            ["E: Медицина", "4", "4", "3", "3", "2", "3", "3", "3", "3", "3", "34/50", "E:21/35 H:10/15"],
        ],
        col_widths=[2.5, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 1.2, 2.5],
        header_bg=BG_GREEN)

    out = os.path.join(DIR, "03_Segments_Scored.docx")
    doc.save(out)
    print(f"  ✓ {out}")


# ════════════════════════════════════════════════════════════════
# 04_Segments_Report.docx
# ════════════════════════════════════════════════════════════════

def generate_04():
    doc = setup()
    title_block(doc,
        "Кому продавать AI-ассистента",
        "Сегментация рынка • RAG Product",
        "Март 2026 • REDO v2 (финальный артефакт)",
        "059669")

    # ─── Summary ───
    h(doc, "Summary", 1, GREEN_DARK)
    p(doc, "Исследованы 9 сегментных гипотез для AI-ассистента с on-prem, методологией внедрения и полным набором функций (документы + email + встречи). 5 сегментов живы, 2 убиты, 2 заморожены/отложены.")
    rich(doc, [
        ("Приоритет 1 — ", {'b': True}),
        ("два параллельных кампейна: ", {}),
        ("IT-компании ", {'b': True}),
        ("(42/50) и ", {}),
        ("Оптовая торговля ", {'b': True}),
        ("(41/50). FL запускает оба одновременно — данные решают какой станет основным.", {}),
    ])

    # ─── Ranking Table ───
    h(doc, "Рейтинг сегментов", 1, GREEN_DARK)
    tbl(doc,
        ["Сегмент", "Приоритет", "Score", "Подтверждено", "Нужна валидация"],
        [
            ["IT-компании", "P1", "42/50", "94 SaaS в рейтинге, CTO в Telegram, ChatGPT заблокирован", "Боль vs «справляемся сами», достаточность функционала"],
            ["Оптовая торговля", "P1 (парал.)", "41/50", "2000-5000 TAM, CRM-gap, 5+ интеграторов", "Тех. валидация КП из прайса, конверсия при 80-110K₽"],
            ["Производство", "P2", "40/50", "НУЛЕВАЯ конкуренция, АСКОН 30 офисов, 3000-5000", "Осознание проблемы, on-prem + PDF/ЕСКД"],
            ["Юр. фирмы", "P2", "39/50", "88% юристов в AI, on-prem (адв. тайна)", "Юр. специфика, реакция партнёров"],
            ["Частная медицина", "P3", "34/50", "120 звонков/день, 80% рутина", "Интеграция с МИС, цена vs альтернативы"],
        ],
        header_bg=BG_GREEN)

    # ─── Killed Segments ───
    h(doc, "Убитые и замороженные сегменты", 1, RED)

    h(doc, "Логистика / Транспорт — УБИТ", 2, RED)
    rich(doc, [("Блокер 1 — Ability to pay: ", {'b': True}), ("При рентабельности 4%, компания 100M₽ = 4M₽ прибыли. Продукт за 960K-1.32M₽/год = 24-33% прибыли. Даже при 500M₽+ = 5-7% прибыли.", {})])
    rich(doc, [("Блокер 2 — ЭДО-триггер ложный: ", {'b': True}), ("Обязательная ТТН обслуживается ЭДО-операторами (1С-ЭПД, Контур). Нулевое пересечение с AI-ассистентом.", {})])
    rich(doc, [("Блокер 3 — AI в логистике ≠ офис: ", {'b': True}), ("45% планируют AI — но маршруты и склады, не офисная рутина.", {})])

    h(doc, "Тренинговые компании — УБИТ", 2, RED)
    rich(doc, [("Блокер 1 — Размер рынка: ", {'b': True}), ("Менее 500 компаний >50M₽. Нет критической массы для FL.", {})])
    rich(doc, [("Блокер 2 — Ценовой fit: ", {'b': True}), ("80-110K₽/мес = 6-8K₽/чел, в 50 раз дороже iSpring LMS (131₽).", {})])

    h(doc, "Госорганизации — ЗАМОРОЖЕН", 2, ORANGE)
    p(doc, "Потенциал подтверждён, но Постановление №1236 запрещает закупку ПО вне Реестра. CorpGPT уже в Реестре. Рекомендация: начать регистрацию немедленно, FL разморозит после получения записи.")

    h(doc, "Девелоперы — ОТЛОЖЕН", 2, ORANGE)
    p(doc, "Сигналы есть (ФСК -45% время документов, КОРТРОС строит базу знаний). Но IT-расходы <1% выручки, конкурент НейроШтат уже на рынке. Мониторить.")

    sep(doc)

    # ════════ IT COMPANIES ════════
    h(doc, "Priority 1: IT-компании (42/50)", 1, BLUE)
    callout(doc,
        "JTBD",
        "Документация устаревает, онбординг долгий, support перегружен, ChatGPT заблокирован в РФ → нужен AI-ассистент для внутренних знаний с on-prem.",
        "2563EB", BG_BLUE, BLUE)

    h(doc, "Профиль компании", 2)
    tbl(doc, ["Параметр", "Значение"],
        [
            ["Отрасль", "Разработка ПО, SaaS, IT-аутсорс (ОКВЭД 62.01, 62.02)"],
            ["Размер", "20-150 сотрудников, выручка 50-300M₽"],
            ["TAM", "300-400 компаний"],
            ["ЛПР", "CTO / VP Engineering (50+ чел.), CEO (<50 чел.)"],
            ["Где активен", "Telegram: @ctodaily, @techdir, @ctoonlive"],
        ], col_widths=[4, 13])

    h(doc, "Как FL достаёт этих людей", 2)
    tbl(doc, ["Канал", "Что делаем", "Результат"],
        [
            ["Email", "firmoteka.ru → email CTO. Hook: «Растёте на 30-50% — как с онбордингом?»", "5-10% reply"],
            ["Telegram DM", "CTO в @ctodaily, @techdir — личное сообщение", "3-8% reply"],
            ["HH.ru", "«technical writer» → компания = наш таргет", "Персонализация"],
            ["Телефон", "Ресепшн → CTO (плоская структура)", "Демо"],
        ], col_widths=[3, 11, 3])

    h(doc, "Топ-5 компаний", 3)
    tbl(doc, ["#", "Компания", "Город", "Выручка", "Сотр.", "Pain signal"],
        [
            ["1", "Cleverence", "Москва", "296M₽", "93", "Сложная документация по мобильным платформам"],
            ["2", "Aspro", "Челябинск", "276M₽", "79", "+50% рост, растущая команда"],
            ["3", "Smartnut", "Екатеринбург", "255M₽", "47", "Борьба с обновлением документации"],
            ["4", "Textback", "Москва", "252M₽", "37", "Быстрый рост, знания не масштабируются"],
            ["5", "Sipuni", "Москва", "243M₽", "99", "+30% рост, support без базы знаний"],
        ], col_widths=[0.8, 2.5, 2.5, 2, 1.2, 7])

    h(doc, "Почему покупают", 2, RED)
    reason(doc, 1,
        "Каждый новый разработчик первые 2 месяца работает вполсилы, потому что не может найти нужную документацию.",
        "При зарплате 200-300K₽/мес вы теряете 200-300K₽ на каждом новом найме. 30% времени на поиск информации [MTS Web Services]. 43% компаний без структурированного онбординга.",
        "AI-ассистент делает нового сотрудника продуктивным с первого дня.")
    reason(doc, 2,
        "Когда уходит senior — вместе с ним уходят знания. Восстановление стоит 2 недели всей команды.",
        "Уход senior = двухнедельный knowledge audit [MTS Web Services]. Знания в чатах, PR, записях встреч — собрать после увольнения невозможно.",
        "AI непрерывно индексирует документы, переписки, записи встреч. Уход человека не обнуляет команду.")
    reason(doc, 3,
        "Support эскалирует тикеты разработчикам. Каждая эскалация = 30 мин разработчика.",
        "20% времени инженеров на поддержку документации [MTS Web Services]. 10 эскалаций/день = 5 часов ≈ 100K₽/мес.",
        "AI ищет по всем источникам разом. Support находит ответ сам.")
    reason(doc, 4,
        "Разработчики пользуются ChatGPT через VPN. Вы не контролируете какой код утекает.",
        "ChatGPT заблокирован в РФ. Команды обходят через VPN на общем аккаунте. Код, переписки, данные клиентов — всё летит в OpenAI.",
        "On-prem AI-ассистент: данные не покидают контур компании.")

    h(doc, "Данные для последующих этапов", 2)
    callout(doc, "Для оффера",
        "ЛПР: CTO / VP Engineering. Идеал: «Новый разработчик продуктивен через 2 недели, не через 2 месяца.» Боль: 7/10. Альтернатива: Confluence (не ищет), ChatGPT VPN (нелегально), кастомная разработка (200K-1M₽). Возражения: «Сделаем сами», «Minervasoft дешевле».",
        "7C3AED", BG_PURPLE, PURPLE)
    callout(doc, "Для GTM",
        "Каналы: Email + Telegram DM + HH.ru + Телефон. Компании: firmoteka.ru (94 SaaS), Rusprofile. Timing: без сезонности, цикл 2-6 недель.",
        "2563EB", BG_BLUE, BLUE)

    h(doc, "План FL", 3)
    p(doc, "50 email + 20 Telegram DM + 10 звонков за 2 недели. Цель: 5-8 ответов, 3-5 discovery calls.", bold=True)

    sep(doc)

    # ════════ WHOLESALE ════════
    h(doc, "Priority 1 (параллельно): Оптовая торговля (41/50)", 1, BLUE)
    callout(doc, "JTBD",
        "Менеджеры по продажам 70% на рутине (КП, переписка, претензии). AI автоматизирует → больше времени на продажи.",
        "2563EB", BG_BLUE, BLUE)

    h(doc, "Профиль компании", 2)
    tbl(doc, ["Параметр", "Значение"],
        [
            ["Отрасль", "Оптовая торговля, дистрибуция (ОКВЭД 46.x)"],
            ["Размер", "30-200 сотрудников, выручка 100-500M₽"],
            ["TAM", "2,000-5,000 компаний по РФ"],
            ["ЛПР", "Коммерческий директор / РОП (50+ чел.), Собственник (<50)"],
            ["Где активен", "Telegram (@optlist_chat 35K, @dirclub 17K), VK"],
        ], col_widths=[4, 13])

    h(doc, "Как FL достаёт этих людей", 2)
    tbl(doc, ["Канал", "Что делаем", "Результат"],
        [
            ["Email", "Rusprofile → email. Hook: «4 менеджера на HH — AI закроет рутину»", "5-8% reply"],
            ["Telegram DM", "@optlist_chat (35K), @dirclub (17K)", "3-5% reply"],
            ["Телефон", "Основной номер → комм. директор. Привыкли к cold calls", "Демо"],
            ["CRM-интеграторы", "CRM Academy, Intervolga, Ингруппа → рекомендации", "Warm referral"],
        ], col_widths=[3, 11, 3])

    h(doc, "Почему покупают", 2, RED)
    reason(doc, 1,
        "Менеджер собирает КП на 47 позиций 40 минут и ошибается в 3 ценах. Клиент уже получил КП от конкурента.",
        "Прайс на 5 000 строк, ручное копирование в Excel. При 10 КП/день = 6+ часов на механику. Ни BitrixGPT ни amoAI не генерируют КП из прайсов — рыночный gap.",
        "AI генерирует КП из прайс-листа за 30 секунд с актуальными ценами.")
    reason(doc, 2,
        "Новый менеджер на 3-й неделе не знает скидку для ключевого клиента. 20 минут на один вопрос.",
        "Условия в CRM, 1С, почте и головах старших. При текучке 30-40% новые менеджеры — постоянная реальность.",
        "AI знает историю по каждому клиенту. Новый менеджер продуктивен с первой недели.")
    reason(doc, 3,
        "РОП проверяет 15 КП в день. В 4 из них ошибки: старые цены, неверная логистика.",
        "Клиент получает КП с прошлогодними ценами — репутационный удар. РОП 3-4 часа/день на проверку.",
        "AI генерирует КП по актуальным данным. Ошибки исключены, РОП высвобожден.")
    reason(doc, 4,
        "Менеджер 2 часа отвечает на претензию — ищет в CRM, почте, 1С. Три системы, ни одна не полная.",
        "5+ претензий/мес = 10+ часов. Медленный ответ = потеря клиента.",
        "AI находит всё за минуту и готовит черновик ответа.")

    callout(doc, "Для оффера",
        "ЛПР: Коммерческий директор / РОП. Идеал: «Менеджеры 70% на продажи вместо 30%.» Боль: 8/10. Альтернатива: CRM + Excel, найм (500-700K₽/мес на 3 чел.). Возражения: «Битрикс и так есть», «Менеджеры не будут пользоваться».",
        "7C3AED", BG_PURPLE, PURPLE)

    p(doc, "План FL: 50 email + 15 Telegram + 15 звонков + 3 CRM-интегратора за 2 недели. Цель: 5-8 ответов, 3-5 discovery calls.", bold=True)

    sep(doc)

    # ════════ MANUFACTURING ════════
    h(doc, "Priority 2: Производство (40/50)", 1, PURPLE)
    callout(doc, "Почему P2, а не P1",
        "Стратегически сильнейший (нулевая конкуренция, АСКОН, TAM 3-5K). Но: осознание проблемы ниже («спроси Михалыча»), outreach медленнее (техдиректора по телефону). FL тестирует параллельно с P1 малым объёмом.",
        "7C3AED", BG_PURPLE, PURPLE)

    tbl(doc, ["Параметр", "Значение"],
        [
            ["JTBD", "Инженеры 60% на нецелевых задачах. Техдокументация разбросана. Уход эксперта = потеря знаний"],
            ["Профиль", "ОКВЭД 25-30, 50-300 сотрудников, 200-1000M₽, КОМПАС-3D, нет PLM/PDM"],
            ["TAM", "3,000-5,000 компаний"],
            ["ЛПР", "Технический директор (рекомендатель) / Гендиректор (бюджет)"],
            ["FL-каналы", "Телефон (primary) + Email + АСКОН-партнёры (30 офисов)"],
        ], col_widths=[3, 14])

    h(doc, "Почему покупают", 2, RED)
    reason(doc, 1,
        "Конструкторы звонят «Михалычу» 4 раза в день. 8 папок на сервере, ни одна не та.",
        "ТУ на сталь 09Г2С в папке «Старое_2019_КОПИЯ». 4 звонка/день × 15 мин = 1 час конструктора + 1 час Михалыча.",
        "AI находит документ за 10 секунд. Михалыч больше не справочное бюро.")
    reason(doc, 2,
        "Михалыч уходит на пенсию — и вместе с ним все знания о допусках и процессах.",
        "Передача: «Если будут вопросы — звоните.» Через полгода звонить некому. Месяцы восстановления.",
        "AI сохраняет знания из документов, переписок, встреч. Уход человека не обнуляет отдел.")
    reason(doc, 3,
        "Новый конструктор 4 месяца разбирается в документации.",
        "Первые 2 месяца не может найти чертёж без помощи. При 80-120K₽/мес: 160-240K₽ потерь.",
        "AI сокращает онбординг: новый инженер находит любой документ с первого дня.")
    reason(doc, 4,
        "ОТК вернул партию: конструктор работал по ТУ 2021 года.",
        "Актуальная версия в другой папке с припиской «_v3_итог_финал». Брак = прямые убытки.",
        "AI всегда отдаёт актуальную версию. Нет путаницы — нет брака.")

    callout(doc, "Ключевой hook",
        "«Аналогов в РФ нет» — единственный отечественный AI для инженерной документации. Это уникальная позиция.",
        "D97706", BG_YELLOW, ORANGE)

    p(doc, "План FL: 20 звонков + 30 email + 3 контакта с АСКОН за 2 недели. Цель: 3-5 discovery calls.", bold=True)

    sep(doc)

    # ════════ LAW FIRMS ════════
    h(doc, "Priority 2: Юр. фирмы (39/50)", 1, PURPLE)
    callout(doc, "Почему P2",
        "Сильнейший use case (88% юристов в AI, on-prem критичен). Но TAM ограничен: ~80-130 компаний. Нишевой beachhead: win → expansion в корп. юр. отделы (2,000-5,000).",
        "7C3AED", BG_PURPLE, PURPLE)

    tbl(doc, ["Параметр", "Значение"],
        [
            ["JTBD", "Экспертный боттлнек. 500+ договоров = 1000-2000 чел.-часов. On-prem (адвокатская тайна)"],
            ["Профиль", "Юр. фирмы (ОКВЭД 69.1), 30-150 юристов, 200-800M₽. Не топ-5 — слишком длинный цикл"],
            ["TAM", "~80-130 → expansion: корп. юр. отделы (2,000-5,000)"],
            ["ЛПР", "Управляющий партнёр. Telegram: «Рульфы» 5K, «ЗарбитражЫ» 1.5K"],
            ["Компании", "КИАП, a.t.Legal, ЦПО Групп, Хренов и Партнёры, Рустам Курмаев"],
        ], col_widths=[3, 14])

    h(doc, "Почему покупают", 2, RED)
    reason(doc, 1,
        "Ассоциаты 3-6 часов ищут прецедент, который лежит в собственном архиве фирмы.",
        "88% юрфирм используют AI, но К+ AI ищет только по своей базе. Прецедент 2019 года — через 2 дня вместо 2 минут.",
        "AI ищет по внутренним документам и внешним базам одновременно — за секунды.")
    reason(doc, 2,
        "500 договоров/год на due diligence. Каждый — 2-3 часа. Это 1000-1500 часов ассоциатов.",
        "86% жалуются на перегрузку. Партнёры перепроверяют — двойная работа. 3-7,5M₽/год на рутину.",
        "AI анализирует договор за минуты, выделяет типовые риски.")
    reason(doc, 3,
        "Юристы не могут использовать ChatGPT — адвокатская тайна. Research полдня вместо 20 минут.",
        "ChatGPT заблокирован + адвокатская тайна = двойной запрет. Junior перебирает 30 вкладок К+ вручную.",
        "On-prem AI: данные не покидают контур фирмы. Compliance + скорость x10.")
    reason(doc, 4,
        "Клиент звонит — помощник 40 минут ищет какая из 4 версий договора финальная.",
        "Без DMS версионирование — хаос. Клиент ждёт.",
        "AI всегда знает актуальную версию. Ответ за секунды.")

    sep(doc)

    # ════════ MEDICINE ════════
    h(doc, "Priority 3: Частная медицина (34/50)", 1, GRAY_600)
    callout(doc, "Самый сложный сегмент для FL",
        "ЛПР не в цифровых каналах. Конкуренты (Chatme.ai) уже на рынке. Цена 80-110K₽ = 11-33% IT-бюджета. Нет партнёров. Тестировать после кейсов из P1/P2.",
        "DC2626", BG_RED, RED)

    tbl(doc, ["Параметр", "Значение"],
        [
            ["JTBD", "120 звонков/день, 80% рутинных. Администраторы перегружены"],
            ["Профиль", "Частные клиники 300M₽-3B₽, 50-300 сотрудников. Не крупные сети"],
            ["TAM", "~150-200 клиник (Vademecum TOP200)"],
            ["ЛПР", "Исполнительный директор / Управляющий клиникой"],
            ["FL-каналы", "Email + Телефон (единственные)"],
            ["Компании", "Здоровье 365, СМТ-Клиника, Парацельс, Клиника Пасман, Авеню"],
        ], col_widths=[3, 14])

    h(doc, "Почему покупают", 2, RED)
    reason(doc, 1,
        "120 звонков/день, 80% рутинные. Администратор 3 минуты на каждый — ищет в бумажном справочнике.",
        "96 рутинных × 3 мин = 4,8 часов. У конкурентов бот за 5 секунд. Пациент на удержании — уходит.",
        "AI отвечает мгновенно, 24/7. Администраторы занимаются сервисом.")
    reason(doc, 2,
        "Текучка администраторов 30% — каждый новый учится 2 месяца.",
        "«Какие анализы перед лапароскопией?» — каждый вопрос эскалируется. Обучение — бесконечный цикл.",
        "AI как мгновенный наставник: новый администратор продуктивен с первого дня.")

    sep(doc)

    # ─── Segment Sequence ───
    h(doc, "Последовательность запуска", 1, GREEN_DARK)
    p(doc, "IT + Торговля (P1, параллельно) → Производство (P2, со сдвигом) → Юр. фирмы (P2) → Медицина (P3)")
    numbered(doc, " Максимально быстрый feedback loop. IT: цикл 2-6 недель, цифровые ЛПР. Торговля: TAM 2-5K, CRM-интеграторы. Оба одновременно.", bold_prefix="IT + Торговля параллельно:")
    numbered(doc, " АСКОН-канал стратегический. Нулевая конкуренция. Но осознание ниже → другой messaging. Начинаем с малого объёма.", bold_prefix="Производство (со сдвигом):")
    numbered(doc, " Малый TAM, но высокая боль. Reference → корп. юр. отделы (10x expansion).", bold_prefix="Юр. фирмы:")
    numbered(doc, " Только после кейсов. Консервативный рынок, конкуренты на месте.", bold_prefix="Медицина:")

    sep(doc)

    # ─── Macro Signal ───
    h(doc, "Макро-сигнал — для всех сегментов", 1, GREEN_DARK)
    callout(doc, "95% AI-пилотов в РФ не дают ROI",
        "75% компаний с бюджетом на AI не имеют стратегии внедрения [MTS Web Services, 700+ компаний]. ChatGPT Pro на 500 сотрудников → через месяц 23 используют. Вывод: методология внедрения — не фича, а ОСНОВНОЙ value prop.",
        "D97706", BG_YELLOW, ORANGE)

    # ─── Recommendations ───
    h(doc, "Рекомендации клиенту", 1, GREEN_DARK)
    numbered(doc, " — блокирует госсектор (2-4 мес., 21% approval rate)", bold_prefix="Начать регистрацию в Реестре ПО")
    numbered(doc, " — критично для Производства и IT", bold_prefix="Подтвердить on-prem deployment")
    numbered(doc, " — ключевой proof point для Торговли", bold_prefix="Подготовить демо КП из реального прайса")
    numbered(doc, " — стратегический канал (30 офисов, 500K рабочих мест)", bold_prefix="Начать диалог с АСКОН")

    sep(doc)

    # ─── Scoring Matrix ───
    h(doc, "Приложение: Сводная матрица оценок", 2)
    tbl(doc,
        ["Сегмент", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9", "C10", "Total"],
        [
            ["IT-компании", "5", "4", "4", "4", "3", "5", "5", "4", "4", "4", "42/50"],
            ["Оптовая торговля", "5", "4", "3", "4", "5", "5", "4", "5", "3", "3", "41/50"],
            ["Производство", "5", "3", "3", "5", "5", "4", "4", "5", "3", "3", "40/50"],
            ["Юр. фирмы", "5", "4", "3", "5", "3", "4", "5", "2", "4", "4", "39/50"],
            ["Медицина", "4", "4", "3", "3", "2", "3", "3", "3", "3", "3", "34/50"],
        ],
        col_widths=[3.5, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1.5],
        header_bg=BG_GREEN)

    # ─── All 9 Hypotheses ───
    h(doc, "Приложение: Все 9 гипотез — финальная судьба", 2)
    tbl(doc,
        ["#", "Гипотеза", "Статус", "Score / Причина"],
        [
            ["1", "Оптовая торговля", "✅ P1", "41/50"],
            ["2", "Производство", "✅ P2", "40/50"],
            ["3", "IT-компании", "✅ P1", "42/50"],
            ["4", "Частная медицина", "✅ P3", "34/50"],
            ["5", "Госорганизации", "❄️ Заморожен", "Реестр ПО — hard blocker"],
            ["6", "Логистика", "❌ УБИТ", "3 блокера (ability to pay, ЭДО, AI≠офис)"],
            ["7", "Сервисные → Юр. фирмы", "✅ P2", "39/50"],
            ["8", "Тренинговые", "❌ УБИТ", "Рынок <500 + цена 50x LMS"],
            ["9", "Девелоперы", "⏸ Отложен", "Рынок не готов, конкурент НейроШтат"],
        ],
        col_widths=[1, 4, 3, 8],
        header_bg=BG_GRAY)

    out = os.path.join(DIR, "04_Segments_Report.docx")
    doc.save(out)
    print(f"  ✓ {out}")


# ════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating DOCX files...")
    generate_02()
    generate_03()
    generate_04()
    print("Done!")
