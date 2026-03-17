#!/usr/bin/env python3
"""Generate beautiful DOCX documents for RAGProduct deliverables."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Color Palette ───
BLUE = RGBColor(0x25, 0x63, 0xEB)
BLUE_DARK = RGBColor(0x1D, 0x4E, 0xD8)
GREEN = RGBColor(0x05, 0x96, 0x69)
GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)
GRAY_600 = RGBColor(0x4B, 0x55, 0x63)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_800 = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Background colors (hex strings for cell shading)
BG_BLUE_LIGHT = "DBEAFE"
BG_GREEN_LIGHT = "D1FAE5"
BG_GREEN_LIGHTER = "ECFDF5"
BG_YELLOW_LIGHT = "FEF3C7"
BG_RED_LIGHT = "FEE2E2"
BG_PURPLE_LIGHT = "F5F3FF"
BG_GRAY_LIGHT = "F3F4F6"
BG_GRAY = "E5E7EB"


# ─── Helpers ───

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, side, color="000000", sz="4", val="single"):
    """Set border on a specific side of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    border = parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
    )
    existing = tcBorders.find(qn(f'w:{side}'))
    if existing is not None:
        tcBorders.remove(existing)
    tcBorders.append(border)


def add_title_block(doc, title, subtitle, version, color_hex):
    """Add a colored title block using a full-width table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color_hex)

    # Remove cell margins / padding
    for side in ['top', 'bottom', 'start', 'end']:
        set_cell_border(cell, side, color=color_hex, sz="0", val="none")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)
    run2.font.color.rgb = WHITE
    run2.font.name = 'Calibri'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(40)
    run3 = p3.add_run(version)
    run3.font.size = Pt(10)
    run3.font.color.rgb = WHITE
    run3.font.name = 'Calibri'

    doc.add_paragraph()  # spacer


def add_heading(doc, text, level=1, color=BLUE):
    """Add a colored heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=None, space_after=None, alignment=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p


def add_rich_para(doc, parts, space_after=None, space_before=None):
    """Add paragraph with mixed formatting. parts = [(text, {bold, italic, color, size}), ...]"""
    p = doc.add_paragraph()
    for text, fmt in parts:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        if fmt.get('bold'):
            run.font.bold = True
        if fmt.get('italic'):
            run.font.italic = True
        if fmt.get('color'):
            run.font.color.rgb = fmt['color']
        if fmt.get('size'):
            run.font.size = Pt(fmt['size'])
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.name = 'Calibri'
    else:
        # Clear auto text and add our own
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.name = 'Calibri'
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
    return p


def add_styled_table(doc, headers, rows, col_widths=None, header_bg=BG_BLUE_LIGHT, highlight_rows=None, highlight_bg=BG_GREEN_LIGHTER):
    """Add a nicely formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        set_cell_shading(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = GRAY_700

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_callout_box(doc, title, text, border_color_hex="2563EB", bg_hex=BG_BLUE_LIGHT, title_color=BLUE):
    """Add a callout box using a 2-column table (colored border + content)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left border cell (narrow colored strip)
    left_cell = table.cell(0, 0)
    left_cell.width = Cm(0.3)
    set_cell_shading(left_cell, border_color_hex)
    left_cell.text = ""

    # Content cell
    right_cell = table.cell(0, 1)
    set_cell_shading(right_cell, bg_hex)

    p = right_cell.paragraphs[0]
    run_t = p.add_run(title)
    run_t.font.bold = True
    run_t.font.color.rgb = title_color
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(11)

    p2 = right_cell.add_paragraph()
    run_c = p2.add_run(text)
    run_c.font.name = 'Calibri'
    run_c.font.size = Pt(10)

    # Remove borders from outer table
    for cell in [left_cell, right_cell]:
        for side in ['top', 'bottom', 'end']:
            set_cell_border(cell, side, val="none", sz="0")
        set_cell_border(cell, 'start', val="none", sz="0")

    doc.add_paragraph()  # spacer
    return table


def add_separator(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def setup_doc():
    """Create a document with default styling."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = GRAY_800

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


# ════════════════════════════════════════════════════════════════
# DOCUMENT 1: POSITIONING
# ════════════════════════════════════════════════════════════════

def generate_positioning():
    doc = setup_doc()
    add_title_block(doc, "AI-ассистент для бизнеса", "Product Positioning Document\nПо методологии April Dunford «Obviously Awesome»", "Version 6.0", "2563EB")

    # ── Product Deep-Dive ──
    add_heading(doc, "Product Deep-Dive", level=1, color=GREEN)
    add_para(doc, "Глубокое понимание продукта — основа позиционирования", italic=True, color=GRAY_600, size=10)

    add_heading(doc, "A. Переформулировка продукта", level=2, color=GRAY_700)
    add_rich_para(doc, [
        ("Это ", {}),
        ("AI-ассистент в чистом виде", {"bold": True}),
        (" — не просто RAG-система для поиска по документам, а полноценный рабочий инструмент, который:", {}),
    ])
    for item in [
        "Подготовит ответ на письмо клиенту или партнёру",
        "Проанализирует записи переговоров и выделит ключевое",
        "Подготовит ответ на претензию с учётом внутренних регламентов",
        "Даст справку по любому вопросу — поищет и в интернете, и во внутренних документах",
        "Суммаризирует длинный документ или переписку",
        "Предложит варианты решения рабочих вопросов",
    ]:
        add_bullet(doc, item)

    add_rich_para(doc, [
        ("Плюс к инструменту: ", {}),
        ("методология внедрения", {"bold": True}),
        (" (как встроить AI в рабочие процессы) и ", {}),
        ("обучение сотрудников", {"bold": True}),
        (" (чтобы реально использовали, а не забросили).", {}),
    ])
    add_rich_para(doc, [
        ("Важно: ", {"bold": True}),
        ("продукт — это ", {}),
        ("лид-магнит", {"bold": True}),
        (", entry point в набор инструментов. Не конечное решение, а начало трансформации.", {}),
    ])

    # B. Transformation Map
    add_heading(doc, "B. Transformation Map", level=2, color=GRAY_700)
    add_callout_box(doc,
        "State A (до):",
        "Сотрудники тратят время на рутину: пишут типовые письма вручную, ищут информацию по папкам и коллегам, готовят ответы на претензии с нуля каждый раз. Компания слышала про AI, хочет «что-то внедрить», но не понимает как начать и боится передавать данные в ChatGPT.",
        "DC2626", BG_RED_LIGHT, RED
    )
    add_callout_box(doc,
        "State B (после):",
        "Сотрудники делают ту же работу, но быстрее: AI готовит черновики писем, находит нужную информацию за секунды, предлагает варианты ответов. Время экономится на рутине, освобождается для сложных задач. Данные при этом не уходят наружу — руководство и СБ спокойны.",
        "059669", BG_GREEN_LIGHT, GREEN
    )

    # C. The Hard Part
    add_heading(doc, "C. The Hard Part — что сложно скопировать", level=2, color=GRAY_700)
    add_callout_box(doc,
        "Методология внедрения из 4 этапов:",
        "1. Анализ типовых бизнес-процессов для конкретного сегмента\n2. Подготовка карты внедрения — какие процессы автоматизировать, в каком порядке\n3. Внедрение типовых инструментов — настройка под конкретную компанию\n4. Поддержка — сопровождение, кастомизация, развитие\n\nЭто не технология (технологию можно скопировать), а экспертиза в применении. Для каждого сегмента — свои типовые процессы, свои боли, свои метрики успеха.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # D. Aha Moment
    add_heading(doc, "D. The «Aha Moment»", level=2, color=GRAY_700)
    add_para(doc, "Типичные триггеры:")
    for item in [
        "Руководитель видит, как конкуренты используют AI, и спрашивает: «А мы почему нет?»",
        "Сотрудники уже используют ChatGPT для работы — но туда нельзя загружать внутренние данные",
        "Пробовали внедрить AI сами — не получилось, сотрудники не используют",
        "Ключевой эксперт перегружен однотипными вопросами от коллег",
    ]:
        add_bullet(doc, item)

    # E. Hidden Value
    add_heading(doc, "E. Hidden Value Propositions", level=2, color=GRAY_700)
    for title, desc in [
        ("Трансформация, не инструмент — ", "результат это оптимизированные бизнес-процессы, а не «ещё один софт»"),
        ("Обучение = adoption — ", "главная причина провала AI-внедрений: сотрудники не используют. Обучение это решает"),
        ("Entry point — ", "начав с AI-ассистента, компания получает партнёра для дальнейшей автоматизации"),
        ("Защита от рисков — ", "не нужно передавать данные наружу, СБ одобрит"),
    ]:
        add_bullet(doc, desc, bold_prefix=title)

    add_separator(doc)

    # ── Executive Summary ──
    add_heading(doc, "Требуется ваше решение", level=1, color=ORANGE)
    add_para(doc, "Позиционирование готово на 85%. Для финализации нужно принять решения:")

    add_styled_table(doc,
        ["Решение", "Статус", "Описание"],
        [
            ["Сегменты исследованы (v2)", "Обновлено ✓", "9 гипотез → 5 живых сегментов. IT-компании и Оптовая торговля — два параллельных приоритета. 50+ компаний в базе."],
            ["Ценообразование", "Отдельный док", "Детальная стратегия монетизации в отдельном документе."],
        ],
        header_bg=BG_YELLOW_LIGHT,
        highlight_rows=[0],
        highlight_bg=BG_GREEN_LIGHTER,
    )

    add_separator(doc)

    # ── Gap Analysis ──
    add_heading(doc, "Gap Analysis", level=1, color=BLUE)
    add_styled_table(doc,
        ["Категория", "Статус", "Комментарий"],
        [
            ["Возможности продукта", "Достаточно ✓", "AI-ассистент: письма, анализ переговоров, справки, суммаризация, ответы на претензии + RAG + гибридный доступ"],
            ["Контекст клиента", "Достаточно ✓", "Компании с БП, которые можно оптимизировать; хотят AI, но не знают как; боятся утечки данных"],
            ["Болевые точки", "Достаточно ✓", "Время на рутину, неэффективные БП, страх утечки данных, «пробовали сами — не получилось»"],
            ["Методология внедрения", "Достаточно ✓", "4 этапа: Анализ БП → Карта внедрения → Внедрение инструментов → Поддержка"],
            ["Сегменты", "Достаточно ✓", "9 гипотез → 5 живых сегментов (2 убиты, 2 заморожены/отложены). Приоритет 1: IT + Торговля (параллельно). 50+ компаний в базе"],
            ["Доказательства", "Частично ⚠", "Есть пилоты, но нужны замеры: время до/после на конкретных задачах"],
            ["Бизнес-модель", "Достаточно ✓", "Внедрение бесплатно, ежемесячный ≤110k, маржа ≥250%. Доход на поддержке, кастомизации, обучении"],
        ],
        highlight_rows=[5],
        highlight_bg=BG_YELLOW_LIGHT,
    )

    add_separator(doc)

    # ── 1. Competitive Alternatives ──
    add_heading(doc, "1. Competitive Alternatives", level=1, color=BLUE)
    add_para(doc, "«Если вашего продукта не существует, что делают клиенты вместо этого?»", italic=True, color=GRAY_600)

    add_callout_box(doc,
        "Контекст: рынок РФ",
        "Западных AI-решений на российском рынке фактически нет. ChatGPT заблокирован, Microsoft Copilot не продаётся, Claude ограничен. Рынок AI-интеграции и трансформации для среднего бизнеса пуст — нет ни одного игрока, который бы его занял. Just AI — единственная российская компания с заметным присутствием, но с другим фокусом.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_styled_table(doc,
        ["Альтернатива", "Что делает адекватно", "Где проигрывает"],
        [
            ["1. «Сами разберёмся»", "Контроль, кастомизация под себя", "Нет экспертизы, месяцы экспериментов, сотрудники не используют. 95% AI-пилотов в РФ не дают ROI"],
            ["2. Ручная работа", "Привычно, понятно, не требует изменений", "Теряется время на рутине, конкуренты уходят вперёд"],
            ["3. ChatGPT через VPN", "Мощная LLM, хороший UX, сотрудники уже используют втихую", "Заблокирован — доступ через VPN (серая зона). Нельзя загружать внутренние документы. СБ не одобрит"],
            ["4. Российские нишевые решения (Minervasoft, AutoFAQ, SalesAI, К+ AI, Doczilla)", "Работают в РФ, решают отдельные задачи", "Каждый закрывает одну функцию. Ни один не покрывает полный набор (документы + email + встречи + on-prem). Нет методологии"],
            ["5. Just AI", "Известный бренд, опыт в разговорном AI", "Фокус на чат-ботах и голосовых ассистентах, не на AI-трансформации БП. Нет методологии для среднего бизнеса"],
        ],
    )

    add_separator(doc)

    # ── 2. Unique Attributes ──
    add_heading(doc, "2. Unique Attributes", level=1, color=BLUE)
    add_para(doc, "«Что есть у вас, чего нет у альтернатив?»", italic=True, color=GRAY_600)

    add_styled_table(doc,
        ["Атрибут", "Описание", "Против кого", "Тип"],
        [
            ["Методология внедрения (4 этапа)", "Анализ БП → Карта → Внедрение → Поддержка. Для каждого сегмента — своя адаптация", "Все альтернативы", "Жёсткий"],
            ["Обучение сотрудников", "Личные тренинги + train-the-trainer модель. Без обучения AI-внедрения проваливаются", "Все альтернативы (ни одна не включает обучение)", "Жёсткий"],
            ["AI-ассистент полного цикла", "Не только RAG: письма, анализ переговоров, претензии, суммаризация, справки", "Чистые RAG-решения", "Мягкий"],
            ["Гибридный доступ к данным", "Интернет + внутренние документы — один инструмент", "ChatGPT (только интернет), закрытые RAG (только внутренние)", "Мягкий"],
            ["Защита данных", "Внутренние документы не передаются наружу + ролевое разграничение доступа", "Публичные AI-сервисы", "Мягкий"],
            ["Сегментная экспертиза", "Опыт внедрения в конкретных вертикалях: торговля, производство, юридические фирмы", "Универсальные решения", "Жёсткий"],
        ],
    )

    # Moat narratives
    add_heading(doc, "Strategic Differentiation Narrative", level=2, color=GRAY_700)

    for title, text in [
        ("Почему это невозможно скопировать?",
         "Технологию (RAG, чат, интеграции) скопировать можно за месяцы. Методологию внедрения — нет. Она требует понимания типовых БП в каждом сегменте, опыта внедрений и навыка обучения сотрудников. Каждое внедрение делает методологию сильнее."),
        ("Почему клиент не вернётся к альтернативам?",
         "Switching costs: операционные (сотрудники обучены), данные (загружена база знаний, настроены доступы), отношения (поддержка и кастомизация создают партнёрство, не подписку)."),
        ("Moat Analysis",
         "• Expertise lock-in: Методология + сегментная экспертиза накапливаются с каждым внедрением\n• Process embedding: AI становится частью рабочих процессов\n• Switching costs: Обученные сотрудники + данные + доступы\n• Network effects: Слабые (пока). Потенциал: база best practices по сегментам"),
    ]:
        add_callout_box(doc, title, text, "2563EB", BG_BLUE_LIGHT, BLUE)

    add_separator(doc)

    # ── 3. Value Map ──
    add_heading(doc, "3. Value Map", level=1, color=BLUE)
    add_para(doc, "«Какую ценность эти уникальные атрибуты создают для клиентов?»", italic=True, color=GRAY_600)

    add_styled_table(doc,
        ["Атрибут", "Что даёт", "Почему важно", "Уверенность"],
        [
            ["Методология внедрения", "Работающий результат, а не «ещё один софт»", "Оптимизированные БП, экономия времени", "Claimed"],
            ["Обучение сотрудников", "Люди реально используют AI", "Главная причина провала — забрасывают через неделю", "Claimed"],
            ["AI-ассистент полного цикла", "Экономия времени на рутине", "Письма, претензии, справки — быстрее", "Claimed"],
            ["Гибридный доступ", "Один инструмент для всего", "Не нужно переключаться между системами", "Inferred"],
            ["Защита данных", "СБ и руководство одобрят", "Снимается барьер «а вдруг утечёт»", "Claimed"],
            ["Сегментная экспертиза", "Понимание специфики бизнеса", "«Они понимают наши процессы»", "Claimed"],
        ],
    )

    add_callout_box(doc,
        "Ключевые метрики ценности",
        "Основная метрика: Время — сколько часов/минут экономится на типовых задачах.\nСопутствующая метрика: Качество — снижение ошибок, повышение консистентности ответов.\n\nЭти метрики нужно замерять в пилотах для перевода value claims из «Claimed» в «Proven».",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_separator(doc)

    # ── 4. Target Segments ──
    add_heading(doc, "4. Target Segments", level=1, color=BLUE)
    add_para(doc, "«Кому эта уникальная ценность важнее всего?»", italic=True, color=GRAY_600)
    add_para(doc, "Исследовано 9 сегментных гипотез. 5 сегментов живы, 2 убиты с доказательствами, 2 заморожены/отложены.")

    add_callout_box(doc,
        "Макросигнал",
        "95% AI-пилотов в РФ не дают ROI (Сколково AI Lab / TAdviser). 75% компаний с бюджетом на AI не имеют стратегии внедрения. Пример: ChatGPT Pro на 500 сотрудников → через месяц 23 используют. Вывод: методология внедрения — не фича, а основная ценность продукта.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_heading(doc, "Приоритеты сегментов", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Сегмент", "Приоритет", "Что подтверждено", "Что нужно проверить"],
        [
            ["IT-компании", "Приоритет 1", "94 SaaS-компании в базе, CTO в Telegram, ChatGPT заблокирован", "Реальная боль vs «справляемся сами»"],
            ["Оптовая торговля", "Приоритет 1", "2000-5000 компаний, CRM-интеграторы как канал, 972+ вакансий", "Генерация КП из прайса, конверсия при цене 80-110K₽"],
            ["Производство", "Приоритет 2", "Нулевая конкуренция, АСКОН (30 офисов) как канал, 3000-5000 компаний", "Осознание проблемы техдиректорами, on-prem + PDF готовность"],
            ["Юридические фирмы", "Приоритет 2", "88% юристов используют AI, on-prem критичен (адв. тайна)", "Юридическая специфика продукта, реакция партнёров"],
            ["Частная медицина", "Приоритет 3", "120 звонков/день, 80% рутина", "Интеграция с мед. системами, ценовая конкуренция"],
        ],
        highlight_rows=[0, 1],
    )

    add_heading(doc, "Последовательность выхода на сегменты", level=2, color=GRAY_700)
    add_callout_box(doc,
        "",
        "Приоритет 1 (параллельно): IT-компании + Оптовая торговля\n  → Приоритет 2 (со сдвигом 1-2 нед.): Производство + Юр. фирмы\n    → Приоритет 3 (после первых кейсов): Частная медицина\n\nIT и Торговля запускаются одновременно — данные за 2 недели определят, какой станет основным. Кейс в IT-компании имеет максимальную ценность: «если IT-компания купила вместо того чтобы делать сама — продукт стоящий».",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # Segment cards
    segments = [
        {
            "badge": "ПРИОРИТЕТ 1",
            "name": "IT-компании",
            "profile": "Разработка ПО, SaaS, IT-аутсорс. 20-150 сотрудников, выручка 50-300M₽. Собственный продукт/платформа, внутренняя документация, support-команда. TAM: 300-400 компаний.",
            "pain": "Инженеры 20% времени на документацию, новые сотрудники 30% на поиск информации. ChatGPT заблокирован в РФ — нет легальной альтернативы. Уход senior-разработчика = двухнедельный аудит знаний.",
            "why_buy": [
                ("Каждый новый разработчик первые 2 месяца работает вполсилы", " — не может найти нужную документацию. При зарплате 200-300K₽/мес вы теряете 200-300K₽ на каждом новом найме. → AI делает нового сотрудника продуктивным с первого дня."),
                ("Когда уходит senior — вместе с ним уходят знания о продукте.", " Восстановление стоит 2 недели всей команды. Знания раскиданы по чатам, PR, записям встреч. → AI непрерывно индексирует все источники. Уход человека не обнуляет команду."),
                ("Support эскалирует тикеты разработчикам", ", потому что не может найти ответ в Confluence. Каждая эскалация = 30 мин разработчика. При 10 эскалациях/день = 5 часов ≈ 100K₽/мес. → AI ищет по всем источникам разом. Support находит ответ сам."),
                ("Разработчики пользуются ChatGPT через VPN", " — вы не контролируете, какой код и данные утекают наружу. → On-prem AI-ассистент: данные не покидают контур компании."),
            ],
            "competition": "Minervasoft (50-150K₽/мес, база знаний без AI-ассистента), AutoFAQ (только support). ChatGPT Team недоступен в РФ. Нет конкурента с полным набором + методология + on-prem.",
            "dmp": "CTO / VP Engineering / технический директор. Активен в Telegram (@ctodaily, @techdir). Цикл сделки 2-6 недель.",
            "companies": "Cleverence (296M₽, сложная документация по платформам), Aspro (276M₽, +50% рост), Smartnut (255M₽, борьба с обновлением документации), Textback (252M₽, знания не масштабируются), Sipuni (243M₽, +30% рост, support нуждается в базе знаний).",
            "fl_channels": "Email: список из firmoteka.ru + HighTime SaaS рейтинг, email CTO через сайт/hunter.io. Telegram: личное сообщение CTO из @ctodaily, @techdir. Телефон: IT-компании плоские, CTO доступен через ресепшн. Сигнал боли: вакансия «technical writer» на HH.ru.",
            "why_priority": "Самый быстрый цикл (2-6 недель), цифровые каналы (CTO в Telegram), максимальная ценность кейса для других сегментов.",
        },
        {
            "badge": "ПРИОРИТЕТ 1 (параллельно)",
            "name": "Оптовая торговля",
            "profile": "Оптовые торговые дома, дистрибьюторы. 30-200 сотрудников, выручка 100-500M₽. Отдел продаж 5+ менеджеров, есть CRM (Битрикс24 / AmoCRM / 1С). TAM: 2000-5000 компаний.",
            "pain": "Менеджеры тратят 70% времени на рутину (КП, переписка, претензии) и только 28-34% на продажи. Стоимость проблемы: ~400K₽/мес на 10 менеджеров. 972+ вакансий на HH.ru = компании нанимают вместо автоматизации.",
            "why_buy": [
                ("Менеджер собирает КП на 47 позиций 40 минут и ошибается в 3 ценах.", " Клиент к этому моменту уже получил КП от конкурента. Прайс на 5 000 строк, ручное копирование. При 10 КП/день = 6+ часов. → AI генерирует КП из прайс-листа за 30 секунд."),
                ("Новый менеджер на 3-й неделе не знает скидку для ключевого клиента.", " 20 минут на один вопрос — звонки старшему, CRM, 1С. При текучке 30-40% новые менеджеры — постоянная реальность. → AI знает историю по каждому клиенту. Новый менеджер продуктивен с первой недели."),
                ("РОП проверяет 15 КП в день. В 4 из них ошибки:", " старые цены, неверная логистика, забыли скидку. Клиент получает КП с прошлогодними ценами — репутационный удар. РОП тратит 3-4 часа/день на проверку. → AI генерирует КП по актуальным данным."),
                ("Менеджер 2 часа отвечает на претензию клиента", " — ищет условия в CRM, переписку в почте, акты в 1С. Три системы, ни одна не даёт полной картины. → AI находит всё за минуту и готовит черновик ответа."),
            ],
            "competition": "BitrixGPT и amoAI не генерируют КП из прайс-листов. SalesAI (39-56K₽, только звонки). Нет конкурента с методологией + AI-ассистент для полного цикла продаж.",
            "dmp": "Коммерческий директор / руководитель отдела продаж. В Telegram ограниченно (@optlist_chat 35K, @dirclub 17K). Основной канал — телефон + email.",
            "companies": "ГК «КСК» (СПб, стройматериалы, нанимают менеджеров для cold sales), ГК KOLOBOX (Н. Новгород, шины/диски, топ-10, 110+ точек), «ГРАТ-ВЕСТ» (Москва, FMCG), «ВЕКТОР» / Total Tools (Москва, инструмент), «СМАРТ-ПРОЕКТ» (Москва, стройматериалы).",
            "fl_channels": "Email: Rusprofile Pro + ExportBase (2459 дистрибьюторов). Телефон: основной канал, торговые привыкли к холодным звонкам. CRM-интеграторы (тёплый канал): CRM Academy (300+ клиентов), Intervolga, Ингруппа. Сигнал боли: вакансия «менеджер по оптовым продажам» = компания растёт.",
            "why_priority": "Крупнейший рынок (2000-5000 компаний), CRM-интеграторы как тёплый канал, готовая боль (70% на рутине).",
        },
        {
            "badge": "ПРИОРИТЕТ 2",
            "name": "Производство",
            "profile": "Обрабатывающее производство. 50-300 сотрудников, 200-1000M₽. Инженерный отдел 3+ чел., используют КОМПАС-3D, нет PLM. TAM: 3000-5000 компаний.",
            "pain": "Инженеры до 60% времени на нецелевых задачах. Техдокументация на файловых серверах + в головах. Онбординг инженера: 3-6 мес. «Спроси Михалыча» — привычная модель, но хрупкая.",
            "why_buy": [
                ("Конструкторы звонят «Михалычу» 4 раза в день", ", потому что не могут найти документ на файловом сервере. 8 папок, ни одна не та. ТУ на сталь 09Г2С — в папке «Старое_2019_КОПИЯ». При 4 звонках/день × 15 мин = 1 час конструктора + 1 час Михалыча. → AI находит нужный документ за 10 секунд."),
                ("Михалыч уходит на пенсию — и вместе с ним уходят все знания", " о допусках, процессах и «так было всегда». Через полгода вопросы начнутся — звонить будет некому. → AI сохраняет знания из документов, переписок, записей встреч. Уход человека не обнуляет отдел."),
                ("Новый конструктор 4 месяца разбирается в документации.", " Первые 2 месяца не может найти нужный чертёж без помощи. При зарплате 80-120K₽/мес: 160-240K₽ потерь на каждом новом инженере. → AI сокращает онбординг: новый инженер находит любой документ с первого дня."),
                ("ОТК вернул партию: конструктор работал по ТУ 2021 года", ", а актуальная версия — в другой папке, с припиской «_v3_итог_финал». Брак из-за устаревшей документации — прямые убытки: материалы, время, репутация. → AI всегда отдаёт актуальную версию документа."),
            ],
            "competition": "Нулевая конкуренция. ЛОЦМАН:PLM (3-8M₽, мы в 4-10x дешевле), TEAMLY (199₽/чел, не для инженерии). AI-конкурентов нет вообще.",
            "dmp": None,
            "companies": None,
            "fl_channels": "Телефон (основной): приёмная завода → техдиректор. АСКОН-партнёры (лучший канал): 30 региональных офисов, каждый знает клиентов КОМПАС в регионе. Email: info@ → персонализированное письмо. Сигнал боли: вакансия «инженер-конструктор КОМПАС-3D» (3146 на HH.ru).",
            "why_priority": "Стратегически сильнейший (нулевая конкуренция, АСКОН как канал). Но осознание проблемы ниже и каналы медленнее. Запуск с малым объёмом параллельно с P1.",
        },
        {
            "badge": "ПРИОРИТЕТ 2",
            "name": "Юридические фирмы",
            "profile": "Юридические фирмы, 30-150 юристов, выручка 200-800M₽. Не топ-5 (ЕПАМ, Б1) — слишком длинный цикл для FL. TAM: ~80-130 компаний (ограничен, но с потенциалом расширения на корпоративные юр. отделы — 2000-5000 компаний).",
            "pain": "Экспертный боттлнек — 500+ договоров = 1000-2000 чел.-часов. 86% юристов с выгоранием. On-prem критичен из-за адвокатской тайны. 88% юристов уже используют AI, но ни один инструмент не работает с внутренними документами.",
            "why_buy": [
                ("Ассоциаты тратят 3-6 часов на поиск прецедента, который лежит в собственном архиве фирмы.", " К+ AI ищет только по своей базе — не по внутренним документам. Прецедент из дела 2019 года находится через 2 дня вместо 2 минут. → AI ищет одновременно по внутренним документам и внешним базам."),
                ("500 договоров в год на due diligence. Каждый — 2-3 часа рутинной проверки.", " Это 1 000-1 500 часов ассоциатов. Партнёры вынуждены перепроверять — двойная работа. При ставке 3-5K₽/час: 3-7,5M₽/год. → AI анализирует договор за минуты, выделяет типовые риски."),
                ("Юристы не могут использовать ChatGPT — адвокатская тайна.", " ChatGPT заблокирован. Даже если обойти — адвокатская тайна запрещает выгружать данные. Junior перебирает формулировки вручную. → On-prem AI-ассистент: данные не покидают контур фирмы. Скорость ×10."),
                ("Клиент звонит — а помощник 40 минут ищет, какая из 4 версий договора финальная.", " Без DMS-системы версионирование — хаос. → AI всегда знает актуальную версию. Ответ клиенту — за секунды."),
            ],
            "competition": "К+ AI (только база Консультант+), Doczilla Pro (2900₽/чел, «недоработанный»), Нейроюрист (50 запросов/мес). Ни один не покрывает: внутренние документы + email + встречи + on-prem.",
            "dmp": None,
            "companies": "КИАП (top-10 по ширине практик, LegalTech-активны), a.t.Legal, ЦПО Групп, Хренов и Партнёры, Рустам Курмаев и Партнёры.",
            "fl_channels": "Email: Право-300 → email партнёра с сайта фирмы. Telegram: «Рульфы, Ильфы и Инхаусы» (5K), «ЗарбитражЫ» (1.5K). Телефон: юр. фирмы привыкли к деловым звонкам. Особенность: малый рынок → каждый контакт ценен, максимальная персонализация.",
            "why_priority": "Высокая боль и готовность платить. Малый рынок, но кейс открывает путь к корпоративным юр. отделам (рынок в 10 раз больше).",
        },
        {
            "badge": "ПРИОРИТЕТ 3",
            "name": "Частная медицина",
            "profile": "Частные клиники, выручка 300M₽-3B₽, 50-300 сотрудников. Не крупные сети (Медси, ЕМС) — слишком длинный цикл для FL. TAM: ~150-200 клиник.",
            "pain": "120 звонков/день, 80% рутинных. Администраторы перегружены. On-prem решает вопросы врачебной тайны (ФСТЭК не нужен для частных клиник).",
            "why_buy": None,
            "competition": None,
            "dmp": None,
            "companies": "Здоровье 365 (Екатеринбург, #38 Vademecum, +38% рост), СМТ-Клиника (Екатеринбург, +20%), Парацельс (+18%), Клиника Пасман (Новосибирск, новичок TOP200), Авеню (Ростов-на-Дону, франшиза).",
            "fl_channels": None,
            "why_priority": "Конкуренты (Chatme.ai) уже имеют кейсы с клиниками из нашего списка. Цена 80-110K₽ = 11-33% IT-бюджета при альтернативах от 950₽/мес. ЛПР не в цифровых каналах. Тестировать после получения кейсов из IT или Торговли.",
        },
    ]

    for seg in segments:
        add_heading(doc, f"[{seg['badge']}]  {seg['name']}", level=3, color=BLUE)
        add_rich_para(doc, [("Профиль: ", {"bold": True}), (seg["profile"], {})])
        add_rich_para(doc, [("Типовая боль: ", {"bold": True}), (seg["pain"], {})])

        if seg.get("why_buy"):
            add_rich_para(doc, [("Почему покупают:", {"bold": True})], space_after=2)
            for bold_part, rest in seg["why_buy"]:
                add_numbered(doc, rest, bold_prefix=bold_part)

        if seg.get("competition"):
            add_rich_para(doc, [("Конкурентный ландшафт: ", {"bold": True}), (seg["competition"], {})])
        if seg.get("dmp"):
            add_rich_para(doc, [("ЛПР: ", {"bold": True}), (seg["dmp"], {})])
        if seg.get("companies"):
            add_rich_para(doc, [("Пример компаний: ", {"bold": True}), (seg["companies"], {})])
        if seg.get("fl_channels"):
            add_rich_para(doc, [("Как FirstLeads выходит: ", {"bold": True}), (seg["fl_channels"], {})])
        add_rich_para(doc, [("Почему этот приоритет: ", {"bold": True}), (seg["why_priority"], {})])

    # Killed / frozen segments
    add_heading(doc, "Отклонённые и замороженные сегменты", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Сегмент", "Статус", "Причина"],
        [
            ["Логистика / Транспорт", "Отклонён", "Три блокера: (1) при рентабельности 4% продукт = 24-33% чистой прибыли — неподъёмно; (2) обязательная э-ТТН обслуживается ЭДО-операторами; (3) AI в логистике — маршруты и склады, не офисная рутина."],
            ["Тренинговые компании", "Отклонён", "Рынок менее 500 компаний с выручкой >50M₽. Продукт за 80-110K₽/мес в 50 раз дороже LMS-альтернатив."],
            ["Госорганизации", "Заморожен", "Потенциал подтверждён (IT-бюджеты 3-15M₽/год), но ПП №1236 запрещает закупку ПО вне Реестра. CorpGPT уже в Реестре. Рекомендация: начать регистрацию немедленно (2-4 мес.)."],
            ["Девелоперы / Строительство", "Отложен", "Сигналы есть, но отрасль наименее цифровизована (IT-расходы <1% выручки). Конкурент НейроШтат уже предлагает AI для стройки. Мониторить."],
        ],
        header_bg=BG_RED_LIGHT,
    )

    add_heading(doc, "Рекомендации клиенту (вне рамок FL)", level=2, color=GRAY_700)
    for item in [
        "Начать регистрацию в Реестре российского ПО — блокирует весь госсектор (2-4 мес., 21% одобрений)",
        "Подтвердить on-prem deployment — критично для Производства (чертежи/PDF) и IT",
        "Подготовить демо-КП из реального прайс-листа — ключевой proof point для Торговли",
        "Проработать интеграцию с АСКОН — стратегический канал для Производства",
    ]:
        add_bullet(doc, item)

    add_separator(doc)

    # ── 5. Market Category ──
    add_heading(doc, "5. Market Category", level=1, color=BLUE)
    add_para(doc, "«Какая рамка восприятия делает вашу ценность очевидной?»", italic=True, color=GRAY_600)

    add_styled_table(doc,
        ["Стратегия", "Название", "Плюсы", "Минусы"],
        [
            ["A. Head-to-Head", "Корпоративный AI-ассистент", "Понятно, есть спрос", "Западные недоступны в РФ, российские нишевые — узкие. Сравнивать не с кем"],
            ["B. Niche", "AI-ассистент для рос. среднего бизнеса", "Чётко отстраивает от enterprise и западных решений", "«Средний бизнес» — размытое понятие"],
            ["C. New Category", "Методология внедрения AI в БП", "Акцент на экспертизе, а не на софте. Оправдывает премиум. Уникально на рынке РФ", "Требует объяснения, новая категория = больше маркетинга"],
        ],
        highlight_rows=[2],
    )

    add_callout_box(doc,
        "Рекомендация: C. New Category — «Методология внедрения AI в бизнес-процессы»",
        "Рынок РФ пустой — можно занять позицию. Ключевое отличие — в методологии и экспертизе, не в софте. Сразу объясняет, почему это не «ещё один ChatGPT». Позиционирует как экспертов/партнёров, а не вендоров.\n\nАльтернатива по сегментам: IT — «AI-инфраструктура для продуктовых команд». Торговля — «AI для продаж». Производство — «AI для инженерных знаний». Юр. фирмы — «AI для юридической практики с on-prem».",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_separator(doc)

    # ── 6. Relevant Trends ──
    add_heading(doc, "6. Relevant Trends", level=1, color=BLUE)
    for trend, impact in [
        ("Рынок РФ пустой", "Нет устоявшихся игроков в AI-интеграции для среднего бизнеса. Окно возможностей для занятия позиции."),
        ("«Все говорят про AI, но никто не знает как»", "Собственники хотят внедрить AI, но не понимают с чего начать — спрос на методологию."),
        ("Сотрудники уже используют ChatGPT", "AI проник снизу — но без контроля и внутренних данных. Нужна легальная альтернатива."),
        ("Страх утечки данных", "Компании запрещают публичные AI — но альтернативы нет. Защита данных = снятие барьера."),
    ]:
        add_rich_para(doc, [(trend + "  —  ", {"bold": True}), (impact, {"color": GRAY_600})])

    add_separator(doc)

    # ── Positioning Statement ──
    add_heading(doc, "Positioning Statement", level=1, color=BLUE)
    add_callout_box(doc,
        "",
        "Для компаний, у которых есть бизнес-процессы с рутинными задачами (письма, справки, ответы, документы), которые хотят внедрить AI, но не знают как начать и не хотят передавать данные наружу,\n\n[Продукт] — это методология внедрения AI в бизнес-процессы: мы анализируем ваши процессы, внедряем AI-ассистента и обучаем сотрудников — чтобы они реально экономили время на рутине.\n\nВ отличие от ChatGPT через VPN (заблокирован, данные уходят наружу), самостоятельных попыток (95% пилотов не дают ROI) и нишевых решений (каждое закрывает одну задачу), [Продукт] даёт работающий результат: оптимизированные процессы, обученных людей и партнёра для развития.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_separator(doc)

    # ── Elevator Pitches ──
    add_heading(doc, "Elevator Pitches", level=1, color=BLUE)

    add_heading(doc, "One-liner (15 слов)", level=3, color=GRAY_600)
    add_para(doc, "«Внедряем AI в ваши бизнес-процессы — чтобы сотрудники делали работу быстрее.»", italic=True, size=12)

    add_heading(doc, "One paragraph (50 слов)", level=3, color=GRAY_600)
    add_para(doc, "«AI-ассистент, который готовит письма, отвечает на вопросы и работает с вашими документами — без утечки данных наружу. Но главное не софт, а методология: мы анализируем ваши процессы, внедряем инструменты и обучаем команду. Результат — сотрудники реально экономят время на рутине.»", italic=True)

    add_heading(doc, "Extended (150 слов)", level=3, color=GRAY_600)
    add_para(doc, "«Компании хотят использовать AI, но сталкиваются с тремя проблемами. Первая: ChatGPT заблокирован, а данные нельзя отправлять наружу. Вторая: пробовали внедрить сами — не получилось, сотрудники не используют (95% AI-пилотов в РФ не дают ROI). Третья: нишевых решений много, но каждое закрывает одну задачу — а нужен полный набор.\n\nМы решаем все три. Наш AI-ассистент работает с вашими данными, не передавая их наружу. Он не просто ищет по документам — готовит письма, анализирует переговоры, отвечает на претензии, даёт справки.\n\nНо главное — это не софт, а методология внедрения. Мы анализируем ваши бизнес-процессы, составляем карту внедрения, настраиваем инструменты и обучаем сотрудников. Результат — оптимизированные процессы и команда, которая реально экономит время.»", italic=True)

    add_separator(doc)

    # ── Assumptions ──
    add_heading(doc, "Допущения (требуют валидации)", level=1, color=BLUE)
    for a in [
        "Методология = ценность — предполагаем, что компании готовы платить за методологию, а не только за софт",
        "Обучение = adoption — предполагаем, что обученные сотрудники действительно будут использовать AI",
        "Время — главная метрика — экономия времени на рутине — основная ценность для клиентов",
        "Лид-магнит работает — entry point ведёт к поддержке, кастомизации и развитию отношений",
    ]:
        add_bullet(doc, a)

    add_separator(doc)

    # ── Decision Cards ──
    add_heading(doc, "Карточки решений", level=1, color=BLUE)

    add_callout_box(doc,
        "Приоритизация сегментов — Обновлено (v2)",
        "Переработано исследование сегментов: 9 гипотез → 5 живых сегментов. IT-компании воскрешены (ранее отклонены ошибочно). Логистика отклонена с доказательствами. Юр. фирмы выделены. 50+ компаний в базе.\n\nПриоритет 1 (параллельно): IT-компании + Оптовая торговля. Приоритет 2: Производство + Юр. фирмы. Приоритет 3: Частная медицина.\n\nОставшийся вопрос: (1) В каких сегментах были пилоты и какие результаты? (2) Готов ли on-prem deployment для IT и Производства? (3) Стоит ли начать регистрацию в Реестре ПО для разблокирования госсектора?",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_callout_box(doc,
        "Ценообразование — отдельный документ",
        "Модель определена: внедрение бесплатно, ежемесячный платёж ≤110k (2 оклада), маржинальность ≥250%. Доход на поддержке, кастомизации, платном обучении, создании RAG. Детальная стратегия — в документе «Стратегия монетизации».",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    # Footer
    add_para(doc, "Product Positioning Document v6.0 | Методология April Dunford «Obviously Awesome»",
             color=GRAY_600, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(BASE_DIR, "01_positioning/positioning-report-v6.docx"))
    print("✓ positioning-report-v6.docx")


# ════════════════════════════════════════════════════════════════
# DOCUMENT 2: MONETIZATION
# ════════════════════════════════════════════════════════════════

def generate_monetization():
    doc = setup_doc()
    add_title_block(doc, "Стратегия монетизации", "AI-ассистент для бизнеса", "v4.0", "7C3AED")

    # ── Executive Summary ──
    add_heading(doc, "Главное за 1 минуту", level=1, color=ORANGE)

    add_rich_para(doc, [
        ("Что продаём: ", {"bold": True}),
        ("не подписку на софт, а трансформацию бизнес-процессов через AI. Проект внедрения — основной доход, подписка — регулярная база.", {}),
    ])
    add_rich_para(doc, [
        ("Как продаём: ", {"bold": True}),
        ("бесплатная демонстрация на данных клиента → платный проект внедрения → ежемесячная подписка → расширение на новые процессы.", {}),
    ])
    add_rich_para(doc, [
        ("Стадия: Pre-traction", {"bold": True, "color": ORANGE}),
        (" — нет подтверждённых метрик. Стандартная тарифная сетка не подходит — нужен бесплатный вход для доказательства ценности.", {}),
    ])

    add_styled_table(doc,
        ["", "Демонстрация", "Проект внедрения ⭐", "Масштабирование"],
        [
            ["Что это", "Бесплатная проба на ваших данных", "AI в 1 процесс / отдел", "Расширение на другие процессы"],
            ["Проект (разово)", "Бесплатно", "от 600 000 ₽", "от 1 200 000 ₽"],
            ["Подписка (/мес)", "—", "80 000 – 110 000 ₽/мес", "120 000 – 200 000 ₽/мес"],
            ["Средний чек", "0 ₽", "~750 000 ₽", "1 500 000+ ₽"],
        ],
        header_bg=BG_PURPLE_LIGHT,
        highlight_rows=[1, 2, 3],
        highlight_bg=BG_GREEN_LIGHTER,
    )

    add_separator(doc)

    # ── 1. Pricing Philosophy ──
    add_heading(doc, "1. Стратегия ценообразования", level=1, color=PURPLE)
    add_para(doc, "Продукт на стадии pre-traction (0 подтверждённых метрик). Стандартная тарифная сетка «Старт / Рост / Корпоративный» не будет работать — покупатель не подпишется на 80 000 ₽/мес без доказательств.")

    add_heading(doc, "Как продавать: философия ценообразования", level=2, color=GRAY_700)

    add_styled_table(doc,
        ["Вариант", "Суть", "Плюсы", "Минусы"],
        [
            ["A. Подписка на софт", "3 тарифа, ежемесячная оплата", "Масштабируемость", "Убивает позиционирование «методология». Сравнивают с ChatGPT Team ($25/чел)."],
            ["B. Трансформация через бесплатный вход ✓", "Бесплатная демо → проект → подписка", "Снимает риск. Доказывает ценность. Оправдывает премиум.", "Ресурсы на бесплатные демо. Длиннее цикл."],
            ["C. Оплата за результат", "Цена привязана к метрикам", "Максимально справедливо", "Невозможно без proof points. Можно внедрить позже."],
        ],
        highlight_rows=[1],
    )

    add_callout_box(doc,
        "Рекомендация: Вариант B",
        "Позиционирование = «методология внедрения AI», не «платформа». Продукт pre-traction — нужен бесплатный вход. Средний чек 750k из переговоров подтверждает: рынок готов платить за проект.",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_heading(doc, "Обязательные механизмы для pre-traction", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Механизм", "Приоритет", "Реализация"],
        [
            ["Бесплатная демонстрация", "Обязательно", "1–5 документов клиента. AI отвечает на реальные вопросы. 1–2 часа."],
            ["Бесплатный пилот", "Обязательно", "Для стратегических клиентов: полное внедрение в обмен на кейс + метрики."],
            ["Скидка за кейс-стади", "По ситуации", "15–20% скидка на проект в обмен на публикуемый кейс. Первые 3–5 кейсов критичны."],
        ],
    )

    add_separator(doc)

    # ── 2. Customer Journey ──
    add_heading(doc, "2. Путь клиента", level=1, color=PURPLE)
    add_rich_para(doc, [
        ("Вы продаёте не тарифный план, а ", {}),
        ("путь от «хотим попробовать AI» до «AI встроен в наши процессы».", {"bold": True}),
    ])

    add_styled_table(doc,
        ["Этап", "Что происходит", "Цена"],
        [
            ["1. Демонстрация", "Загружаем 1–5 документов клиента, показываем AI на реальных данных. «Aha-moment»: AI отвечает по ЕГО документам.", "Бесплатно"],
            ["1.5. Пилот (опц.)", "Для стратегических клиентов: бесплатное внедрение в 1 процесс, 2–4 недели.", "Бесплатно (инвестиция)"],
            ["2. Проект внедрения ⭐", "Полная методология: аудит → карта → настройка AI → загрузка данных → обучение. 20–160 часов.", "от 600 000 ₽ (средний ~750k)"],
            ["3. Подписка", "AI-ассистент + хостинг + API + поддержка до 10 ч/мес + обновления. До 50 пользователей.", "80 000 – 110 000 ₽/мес"],
            ["4. Расширение", "Новые процессы, отделы, интеграции, дополнительные RAG-системы.", "от 300 000 ₽ + рост подписки"],
        ],
        highlight_rows=[2],
    )

    add_separator(doc)

    # ── 3. Revenue Model ──
    add_heading(doc, "3. Модель поставки и дохода", level=1, color=PURPLE)

    add_callout_box(doc,
        "Модель",
        "Поставка: Гибридная — софт + экспертиза внедрения.\nДоход: Проект (основной) + подписка (регулярный) + расширение (рост).\nМетрика: За компанию/проект (не за пользователя — стимулирует adoption).",
        "7C3AED", BG_PURPLE_LIGHT, PURPLE
    )

    add_heading(doc, "Прогноз на 1 клиента за год", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Источник", "Когда", "Сумма", "Маржа"],
        [
            ["Проект внедрения", "Месяц 0", "600 000 – 1 500 000 ₽", "~70%"],
            ["Подписка (12 мес)", "Месяцы 1–12", "960 000 – 1 320 000 ₽", "~65–75%"],
            ["Расширение", "Месяцы 6–12", "300 000 – 1 000 000 ₽", "~70%"],
            ["Итого за 1-й год", "", "1 860 000 – 3 820 000 ₽", ""],
        ],
        highlight_rows=[3],
    )

    add_separator(doc)

    # ── 4. Price Architecture ──
    add_heading(doc, "4. Ценовая архитектура", level=1, color=PURPLE)

    add_heading(doc, "Проект внедрения", level=2, color=GRAY_700)
    add_para(doc, "Формула: Часы × Ставка (4 000 ₽/час) × Коэффициент маржи (3.5)\nМинимальная цена: 600 000 ₽ (якорь из переговоров)", size=10)

    add_styled_table(doc,
        ["Масштаб проекта", "Часы", "Себестоимость", "Цена (×3.5)"],
        [
            ["Базовый — 1 процесс, до 500 документов", "20–40 ч", "80–160k ₽", "600 000 ₽ (мин.)"],
            ["Стандартный — 2–3 процесса, до 2 000 док.", "40–80 ч", "160–320k ₽", "600 000 – 1 120 000 ₽"],
            ["Расширенный — 5+ процессов, 5 000+ док.", "80–160 ч", "320–640k ₽", "1 120 000 – 2 240 000 ₽"],
        ],
        highlight_rows=[1],
    )

    add_callout_box(doc,
        "Проверка по среднему чеку",
        "Средний чек из переговоров: 750 000 ₽. Стандартный проект (40–60 часов) попадает в этот диапазон — модель объясняет реальную готовность платить.",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_heading(doc, "Ежемесячная подписка", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["", "При 5 клиентах", "При 10 клиентах", "При 20 клиентах"],
        [
            ["Себестоимость на клиента", "~31 000 ₽", "~25 400 ₽", "~22 700 ₽"],
            ["Подписка 100 000 ₽/мес", "Маржа 69%", "Маржа 75%", "Маржа 77%"],
        ],
        highlight_rows=[1],
    )

    add_callout_box(doc,
        "Рекомендация: 100 000 ₽/мес как стартовая точка",
        "100k/мес на фоне проекта 750k — это 13% от первоначальной инвестиции. Выглядит как мелочь. Но за год это 1.2M — больше, чем сам проект. Стратегия «менее заметно, более маржинально».",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_separator(doc)

    # ── 5. Assumptions & Risks ──
    add_heading(doc, "5. Допущения и риски", level=1, color=PURPLE)

    add_styled_table(doc,
        ["Допущение", "Риск если неверно", "Как проверить"],
        [
            ["Бесплатная демо конвертирует в проект", "Бесплатная работа впустую", "Конверсия первых 10 демо. Ориентир: 20–30%"],
            ["Клиенты готовы платить 600k+ за внедрение", "Цена отсекает рынок", "Уже подтверждено: средний чек 750k"],
            ["Подписка 80–110k/мес не вызывает отторжения", "Клиенты уходят после внедрения", "Мин. 12 мес в договоре. Retention — метрика"],
            ["Инфраструктура $600/мес достаточна для 5–10 клиентов", "Затраты > выручки", "Мониторинг нагрузки. >$150/кл → пересчитать"],
            ["Поддержка ≤ 10 ч/мес", "Поддержка съедает маржу", "Учёт времени. >15ч → ввести лимит"],
        ],
    )

    add_callout_box(doc,
        "Валютный риск",
        "Инфраструктура ($600/мес) привязана к доллару. При курсе 90 ₽/$ — 54 000 ₽/мес, при 120 ₽/$ — 72 000 ₽/мес. Включите в договор пункт о пересмотре цены при изменении курса более чем на 15%.",
        "D97706", BG_YELLOW_LIGHT, ORANGE
    )

    add_separator(doc)

    # ── Appendix 1 ──
    add_heading(doc, "Приложение 1: Данные из позиционирования", level=1, color=GRAY_600)

    add_heading(doc, "Сигналы продукта", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Параметр", "Значение"],
        [
            ["Тип продукта", "AI-ассистент полного цикла (RAG + генерация + гибридный доступ)"],
            ["Категория", "Методология внедрения AI в бизнес-процессы (новая категория)"],
            ["Сервисная составляющая", "Значительная: 4-этапная методология + обучение сотрудников"],
            ["Позиционирование", "Премиум-эксперты, не вендоры. Партнёрство, не подписка."],
        ],
        header_bg=BG_GRAY_LIGHT,
    )

    add_heading(doc, "Сигналы клиента", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Параметр", "Значение"],
        [
            ["Основной сегмент", "Торговые и производственные компании, 30–300 сотрудников"],
            ["Покупатель", "Собственник / CEO / HRD"],
            ["Триггер покупки", "«Конкуренты используют AI», «пробовали сами — не получилось», «боимся утечки данных»"],
            ["Средний чек", "750 000 ₽"],
        ],
        header_bg=BG_GRAY_LIGHT,
    )

    add_callout_box(doc,
        "0 из 6 ценностей подтверждены метриками",
        "Это определило стратегию: бесплатный вход + проектная модель. Как только появятся proof points, можно убрать пилоты, поднять цены на 20–30% и добавить «оплата за результат».",
        "D97706", BG_YELLOW_LIGHT, ORANGE
    )

    add_separator(doc)

    # ── Appendix 2 ──
    add_heading(doc, "Приложение 2: Детальные расчёты", level=1, color=GRAY_600)

    add_heading(doc, "Исходные данные", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Параметр", "Значение"],
        [
            ["Хостинг + API", "$600/мес (общий)"],
            ["Почасовая ставка", "4 000 ₽/час"],
            ["Часы внедрения", "20–160 ч"],
            ["Целевая маржа", "от 250% (×3.5)"],
            ["Поддержка", "до 10 ч/мес"],
            ["Средний чек", "750 000 ₽"],
            ["Активные пользователи", "~30 чел"],
            ["Курс доллара", "~90 ₽/$"],
        ],
        header_bg=BG_GRAY_LIGHT,
    )

    add_heading(doc, "Ценовой потолок (EVC)", level=2, color=GRAY_700)
    add_para(doc, "Цена альтернативы (ChatGPT Team для 30 чел): ~3 000 ₽/чел/мес × 30 = 90 000 ₽/мес\nДифференцированная ценность (Claimed, ×0.7): +234 500 ₽/мес\nПотолок (EVC) = ~325 000 ₽/мес", size=10)

    add_heading(doc, "Якоря конкурентов", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Альтернатива", "Их цена", "Наша позиция"],
        [
            ["ChatGPT Team (30 чел)", "~90 000 ₽/мес", "Дороже, но с методологией и защитой данных"],
            ["Заказная разработка RAG", "1.5–5M ₽", "Значительно дешевле + включает методологию"],
            ["Консалтинг по AI (Big4)", "3–10M ₽", "На порядок дешевле при сопоставимом результате"],
            ["«Сами разберёмся»", "0 ₽ + 3–6 мес + 80% провал", "750k — страховка от потери времени"],
        ],
    )

    add_callout_box(doc,
        "Позиция: между ChatGPT и заказной разработкой",
        "Не дёшево (иначе не поверят), не дорого (доступно среднему бизнесу). Правильная позиция для «методологии внедрения AI».",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    # Footer
    add_para(doc, "Стратегия монетизации v4.0 | На основе позиционирования + реальных данных клиента",
             color=GRAY_600, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(BASE_DIR, "02_monetization/monetization-strategy-v4.docx"))
    print("✓ monetization-strategy-v4.docx")


# ════════════════════════════════════════════════════════════════
# DOCUMENT 3: SEGMENTATION
# ════════════════════════════════════════════════════════════════

def generate_segmentation():
    doc = setup_doc()
    add_title_block(doc, "Кому продавать AI-ассистента", "Сегментация рынка и карта первых клиентов", "Version 1.0", "059669")

    # ── Executive Summary ──
    add_heading(doc, "Коротко: что сделали и что поняли", level=1, color=GREEN)

    add_para(doc, "Мы начали с 8 гипотез о том, кому можно продавать AI-ассистента. Провели исследование рынка, оценили каждый сегмент по 10 критериям, нашли конкретные компании. Вот что получилось:")

    # Summary grid as table
    add_styled_table(doc,
        ["Показатель", "Значение"],
        [
            ["Живых сегментов", "4 (из 8 гипотез)"],
            ["Конкретных компаний найдено", "41"],
            ["Именных ЛПР (с должностями)", "10"],
            ["Сегментов убито (с обоснованием)", "3"],
        ],
        header_bg="D1FAE5",
    )

    add_heading(doc, "Рекомендуемая последовательность входа", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Этап", "Сегмент", "Деталь"],
        [
            ["① Начинаем", "Оптовая торговля / Дистрибьюторы", "12 компаний. Цикл 1-3 мес. Самый быстрый путь к первым кейсам."],
            ["② Следующий", "Производство с инженерным офисом", "12 компаний. Цикл 2-4 мес. Ноль AI-конкурентов."],
            ["③ Потом", "Логистика / Транспорт", "9 компаний. Обязательный ЭДО с сентября 2026."],
            ["④ На будущее", "Юридические / Бухгалтерские фирмы", "8 компаний. Боль сильная, но продукт пока дорогой."],
        ],
        header_bg="D1FAE5",
    )

    add_callout_box(doc,
        "Что нужно от вас (30 минут)",
        "1. В каких из этих сегментов уже были пилоты?\n2. К каким руководителям вы можете дотянуться лично?\n3. Какой сегмент интереснее стратегически?\n\nОтветы определяют, с какого сегмента реально начинать. Личный доступ к ЛПР — это то, что нельзя «докупить» маркетингом.",
        "2563EB", BG_BLUE_LIGHT, BLUE
    )

    add_separator(doc)

    # ── How we researched ──
    add_heading(doc, "Как исследовали", level=1, color=BLUE)
    add_para(doc, "Работали в три фазы:")
    add_numbered(doc, " ", bold_prefix="Гипотезы:")
    add_para(doc, "   Из позиционирования вытащили 8 потенциальных сегментов")
    add_numbered(doc, " ", bold_prefix="Скоринг:")
    add_para(doc, "   Каждый сегмент оценили по 10 критериям. Три сегмента убили — там точно не взлетит")
    add_numbered(doc, " ", bold_prefix="Компании:")
    add_para(doc, "   В четырёх живых сегментах нашли конкретные компании с сигналами боли")

    add_heading(doc, "Сводная таблица оценок", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Сегмент", "Балл", "Статус", "Почему такой вердикт"],
        [
            ["Оптовая торговля", "40/50", "Tier 1 ✓", "Самые богатые каналы, быстрый цикл, понятный ROI"],
            ["Производство", "40/50", "Tier 1 ✓", "Ноль AI-конкурентов, наша цена в 4-10 раз ниже PLM"],
            ["Логистика", "33/50", "Tier 2", "Регуляторный дедлайн создаёт срочность, но маржа падает"],
            ["Сервисные (юр/бух)", "33/50", "Tier 2", "Боль сильная, но продукт стоит 2.5-7x IT-бюджета"],
            ["IT-компании", "31/50", "Убит ✕", "Сами соберут RAG дешевле"],
            ["Тренинговые", "29/50", "Убит ✕", "80-110K/мес нереально для 10-15 человек"],
            ["Госорганизации", "29/50", "На потом", "Нужны Реестр ПО + ФСТЭК + 44-ФЗ"],
            ["Медицина", "28/50", "Убит ✕", "ФЗ-152, врачебная тайна, регуляторика"],
        ],
        highlight_rows=[0, 1],
    )

    add_separator(doc)

    # ── Segment 1: Trade ──
    add_heading(doc, "BEACHHEAD: Оптовая торговля / Дистрибьюторы", level=1, color=BLUE)
    add_para(doc, "Торговые дома и дистрибьюторы — самый понятный и быстрый сегмент для старта. Менеджеры по продажам тратят 70% времени на рутину и только 2 часа 23 минуты из 8-часового дня — на клиентов.")

    add_styled_table(doc,
        ["Параметр", "Значение"],
        [
            ["Кого ищем", "Оптовые компании, дистрибьюторы (ОКВЭД 46.x). 30-200 сотрудников, выручка 100-500 млн. Отдел продаж от 5 менеджеров, есть CRM."],
            ["Размер рынка", "5 000 — 15 000 целевых компаний в России. Рынок растёт (+17% в 2024)."],
            ["Кто решает", "Коммерческий директор или руководитель отдела продаж. В компаниях до 50 чел. — собственник."],
            ["Цикл сделки", "1-3 месяца. Proof-of-value за дни."],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_heading(doc, "Главная боль", level=2, color=GRAY_700)
    add_para(doc, "Менеджер тратит 40 минут на подготовку КП из прайс-листа. Потом ещё 20 минут на ответ на претензию. Это повторяется каждый день. 10 менеджеров теряют ~400 000 руб./мес продуктивности на рутине.")

    add_heading(doc, "С кем конкурируем", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Конкурент", "Что делает", "Наше преимущество"],
        [
            ["Bitrix24 CoPilot", "AI-фичи бесплатно внутри CRM", "Мы работаем с внутренними документами + методология + обучение"],
            ["SalesAI", "Речевая аналитика звонков. 190K + 10₽/мин", "Мы решаем шире: КП, переписку, претензии, документы"],
            ["AmoAI", "Базовый AI в AmoCRM", "Фича в CRM, а не целостное решение"],
        ],
    )

    add_callout_box(doc,
        "Риск: CRM-вендоры добавляют AI бесплатно",
        "Bitrix24 и AmoCRM наращивают AI-функции. Если CoPilot научится генерировать КП из прайс-листа — наше преимущество ослабнет. Важно набрать 5-8 клиентов быстро, пока окно открыто.",
        "D97706", BG_YELLOW_LIGHT, ORANGE
    )

    add_heading(doc, "Конкретные компании (топ из 12 найденных)", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Компания", "Город", "Детали", "ЛПР"],
        [
            ["Платан-Энерго", "Москва", "Дистрибьютор электронных компонентов с 1991. Выручка 641 млн. 4 вакансии одновременно.", "Трубицын Дмитрий Сергеевич — ген. директор"],
            ["ОМНИМИКС", "Москва", "Выручка 855 млн (+12%), прибыль +66%. Масштабируется — отдел продаж не успевает.", "Крапивницкий Дмитрий Викторович"],
            ["Русич", "Тамбов", "FMCG-дистрибьютор, 150+ сотрудников. Огромный объём обработки заказов.", "—"],
            ["Апорт", "Москва", "B2B оптовые продажи. Вакансия с требованиями Битрикс24 + 1С + Excel.", "—"],
        ],
    )
    add_para(doc, "+ ещё 8 компаний: РБМ-СТРОЙ, Техноуспех, МедБаза, НВ-Лаб, Красные Крыши, ЦентрМеталлокровли, Компонента, Авалон.", size=9, color=GRAY_600)

    add_heading(doc, "Как добраться", level=2, color=GRAY_700)
    add_rich_para(doc, [("Telegram: ", {"bold": True}), ("@grebenukm (260K подписчиков), @dirclub (17K), @Salesnotes (15K)", {})])
    add_rich_para(doc, [("Партнёры: ", {"bold": True}), ("CRM Academy — платиновый Битрикс24, #1 Агентство года (Ruward 2025), 300+ клиентов", {})])

    add_separator(doc)

    # ── Segment 2: Manufacturing ──
    add_heading(doc, "TIER 1: Производство с инженерным офисом", level=1, color=BLUE)
    add_para(doc, "Самый «чистый» сегмент: ни одного AI-конкурента для инженерных знаний на российском рынке. Ближайшая альтернатива (PLM-система ЛОЦМАН) стоит в 4-10 раз дороже нас.")

    add_styled_table(doc,
        ["Параметр", "Значение"],
        [
            ["Кого ищем", "Обрабатывающее производство (ОКВЭД 25-30). 50-300 сотрудников, 200-1000 млн. Инженерный отдел, КОМПАС-3D, НЕТ PLM."],
            ["Размер рынка", "8 000 — 12 000 компаний. 75-85% без PLM = 6 000 — 10 000 адресуемых."],
            ["Кто решает", "Технический директор рекомендует, ген. директор / собственник утверждает."],
            ["Цикл сделки", "2-4 месяца (двухступенчатый). Дольше, но крупнее средний чек."],
        ],
        header_bg=BG_BLUE_LIGHT,
    )

    add_heading(doc, "Главная боль", level=2, color=GRAY_700)
    add_para(doc, "Инженеры тратят до 60% времени на нецелевые задачи: ищут документацию по файловым серверам, спрашивают у «Михалыча». Онбординг нового инженера — 3-6 месяцев. Если «Михалыч» уходит на пенсию — знания уходят вместе с ним.")

    add_heading(doc, "С кем конкурируем", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Конкурент", "Что делает", "Наше преимущество"],
        [
            ["ЛОЦМАН:PLM", "Управление техдокументацией. 3-8 млн внедрение", "В 4-10 раз дешевле. Быстрый запуск. AI-поиск, а не просто хранилище"],
            ["TEAMLY", "Общекорпоративная база знаний. 199₽/чел", "Для офиса, не для инженерии. Не понимает ТУ, ЕСКД"],
        ],
    )

    add_heading(doc, "Конкретные компании (топ из 12)", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Компания", "Город", "Детали", "ЛПР"],
        [
            ["АО «ЗБО»", "Оренбург", "Буровое оборудование. 895 млн, 247 чел., 328 наименований. Клиенты: Алроса, Норникель. Подтверждённый КОМПАС-3D.", "Медведев А.К. — ген. директор"],
            ["Пожарные системы", "Тверь", "Спецтехника, КОМПАС V16, бета-тестеры АСКОН.", "—"],
            ["НПЦ «ИНФОТРАНС»", "Самара", "Приборостроение: информационные и транспортные системы. КОМПАС V15.", "—"],
        ],
    )
    add_para(doc, "+ Екатеринбургский кластер (5 компаний), КРМЗ, «25 Микрон», ГК «СТАЛТ», «Комлед».", size=9, color=GRAY_600)

    add_heading(doc, "Как добраться", level=2, color=GRAY_700)
    add_rich_para(doc, [("Партнёры: ", {"bold": True}), ("АСКОН — 30 региональных офисов, 500 000 рабочих мест КОМПАС-3D. Открытая партнёрская программа.", {})])
    add_rich_para(doc, [("Мероприятия: ", {"bold": True}), ("ИННОПРОМ — 6-9 июля, Екатеринбург, 52K+ посетителей.", {})])

    add_callout_box(doc,
        "Нюанс: может потребоваться on-prem",
        "Конструкторская документация — коммерческая тайна. Некоторые производства не отдадут данные в облако. Нужно уточнить техническую готовность к on-prem deployment.",
        "D97706", BG_YELLOW_LIGHT, ORANGE
    )

    add_separator(doc)

    # ── Segment 3: Logistics ──
    add_heading(doc, "TIER 2: Логистика / Транспорт", level=1, color=ORANGE)
    add_para(doc, "Логистика переживает три регуляторных шока одновременно: обязательный реестр экспедиторов (ГосЛог) с 1 марта 2026, обязательный ЭДО с 1 сентября 2026, и дефицит кадров. При рентабельности 4% — автоматизация рутины становится вопросом выживания.")

    add_styled_table(doc,
        ["Параметр", "Значение"],
        [
            ["Кого ищем", "Грузоперевозки, экспедирование (ОКВЭД 49.4/52.2). 30-200 сотрудников, 100-500 млн."],
            ["Размер рынка", "5 000 — 8 000 компаний. AI для адм. задач — полностью пустая ниша."],
        ],
        header_bg=BG_YELLOW_LIGHT,
    )

    add_heading(doc, "Компания-идеал", level=2, color=GRAY_700)
    add_styled_table(doc,
        ["Компания", "Город", "Детали", "ЛПР"],
        [
            ["ТК Грузоперевозки", "Москва", "Идеальное попадание. 328 млн (+37% за год), 78 чел. (+26%). FTL-перевозки.", "Бочков С.Н. — директор (с 2015)"],
        ],
    )
    add_para(doc, "+ ещё 8 компаний: ТД Логистик (27 вакансий!), Автологистика СПб, ТК «Открытие», ТИЭСВАЙ, ТК «Подорожник», ТК «ЭКСИС», ПТК1, Альянс Рус Регион.", size=9, color=GRAY_600)

    add_callout_box(doc,
        "TransRussia — 17-19 марта 2026, Крокус Экспо",
        "500+ экспонентов, 30 000+ посетителей. Регистрация бесплатная на transrussia.ru. Через 25 дней.",
        "059669", BG_GREEN_LIGHTER, GREEN
    )

    add_separator(doc)

    # ── Segment 4: Service ──
    add_heading(doc, "НА БУДУЩЕЕ: Юридические и бухгалтерские фирмы", level=1, color=GRAY_600)
    add_para(doc, "Боль здесь одна из самых острых: старшие партнёры тратят 2-3 часа в день на одинаковые вопросы от джуниоров. 86% юристов испытывают выгорание. Но есть блокер.")

    add_callout_box(doc,
        "Ценовой блокер",
        "IT-бюджет юридической фирмы из 20-100 человек — 300-800 тыс. руб./год. Наш первый год (~2 млн) — это 2.5-7 раз больше их всего IT-бюджета. Пока не создадим облегчённый пакет за 40-50K/мес — этот сегмент недоступен.",
        "DC2626", BG_RED_LIGHT, RED
    )

    add_styled_table(doc,
        ["Компания", "Город", "Детали", "ЛПР"],
        [
            ["Forward Legal", "Москва", "50+ юристов (удвоение за 4 года). Свой цифровой сервис «Банкротный таймлайн».", "Эльмира Кондратьева — управляющий партнёр"],
            ["Рустам Курмаев и партнёры", "Москва", "Forbes ТОП-5 юрфирм. ~37 юристов, 2-е место по выручке на юриста.", "Рустам Мусаевич Курмаев"],
        ],
    )
    add_para(doc, "+ ещё 6 компаний: КИАП, Суррей, Хренов и Партнёры, Бетерра, Acsour, ГЛАВБУХ АССИСТЕНТ.", size=9, color=GRAY_600)

    add_separator(doc)

    # ── Events ──
    add_heading(doc, "Календарь мероприятий: первое полугодие 2026", level=1, color=BLUE)
    add_para(doc, "Мероприятия — основной способ выйти на ЛПР в первые месяцы.")

    add_styled_table(doc,
        ["Дата", "Мероприятие", "Город", "Детали"],
        [
            ["12 марта", "Business Force Forum", "Москва", "100+ спикеров, треки AI Force + Sales Force. Кросс-сегмент."],
            ["17-19 марта", "TransRussia 2026", "Крокус Экспо", "500+ экспонентов, 30K+ посетителей. Бесплатная регистрация. Логистика."],
            ["27 марта", "Global Tech Forum", "Москва", "2500+ участников. Цифровая трансформация. Кросс-сегмент."],
            ["19-22 мая", "ЦИПР 2026", "Нижний Новгород", "15K+ участников из 40 стран. Кросс-сегмент."],
            ["24-26 июня", "ПМЮФ 2026", "СПб", "5 600+ участников из 80 стран. Юридический сегмент."],
            ["6-9 июля", "ИННОПРОМ 2026", "Екатеринбург", "52K+ посетителей, 900+ экспонентов. Производство."],
        ],
    )

    add_separator(doc)

    # ── Partners ──
    add_heading(doc, "Партнёры-интеграторы: кто поможет продавать", level=1, color=BLUE)
    add_styled_table(doc,
        ["Партнёр", "Для сегмента", "Почему важен"],
        [
            ["CRM Academy", "Торговля", "Платиновый Битрикс24, #1 Агентство года (Ruward 2025), 300+ клиентов."],
            ["АСКОН", "Производство", "30 региональных офисов, 500K рабочих мест КОМПАС-3D. Открытая партнёрская программа."],
            ["1С-Рарус", "Кросс-сегмент", "#1 франчайзи 1С, 17 офисов, 230 городов."],
            ["Первый Бит", "Кросс-сегмент", "100+ офисов в 90 городах. Платиновый и 1С, и Битрикс24."],
        ],
    )

    add_separator(doc)

    # ── What to do next ──
    add_heading(doc, "Что делать дальше", level=1, color=GREEN)
    add_para(doc, "Три действия, которые дадут максимальный эффект в ближайшие 30 дней:")

    for title, desc in [
        ("1. Определить beachhead (нужен ваш input)", "Торговля или Производство — оба набрали одинаковый балл. Выбор зависит от того, к кому из ЛПР вы можете дотянуться лично."),
        ("2. Зарегистрироваться на ближайшие мероприятия", "Business Force Forum (12 марта) и TransRussia (17-19 марта, бесплатно) — через 20 и 25 дней."),
        ("3. Инициировать контакт с партнёрами", "CRM Academy (crmacademy.ru) для торговли и АСКОН (ascon.ru/for-partners/) для производства."),
    ]:
        add_callout_box(doc, title, desc, "059669", BG_GREEN_LIGHTER, GREEN)

    add_heading(doc, "Первые 4 компании для контакта", level=2, color=GREEN)
    add_styled_table(doc,
        ["#", "Компания", "Почему"],
        [
            ["1", "Платан-Энерго", "Дистрибьютор компонентов, 641 млн, 4 вакансии, ЛПР известен"],
            ["2", "АО «ЗБО»", "Буровое оборудование, 895 млн, подтверждённый КОМПАС без PLM, ЛПР известен"],
            ["3", "ТК Грузоперевозки", "Идеальное попадание, 328 млн, +37% рост, ЛПР известен"],
            ["4", "Forward Legal", "Удвоение команды за 4 года + собственный цифровой продукт, ЛПР известен"],
        ],
        header_bg="D1FAE5",
    )

    # Footer
    add_para(doc, "Сегментация рынка v1.0 | AI-ассистент для бизнеса",
             color=GRAY_600, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(BASE_DIR, "04_segments/segment-research-v1.docx"))
    print("✓ segment-research-v1.docx")


# ════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

if __name__ == "__main__":
    print("Generating DOCX documents...")
    generate_positioning()
    generate_monetization()
    generate_segmentation()
    print("\nDone! All 3 documents generated.")
